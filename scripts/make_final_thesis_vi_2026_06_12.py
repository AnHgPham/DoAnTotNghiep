from __future__ import annotations

import csv
from dataclasses import dataclass
from pathlib import Path
from typing import Iterable

import matplotlib.pyplot as plt
import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "docs" / "thesis"
ASSET_DIR = OUT_DIR / "assets_final_2026_06_12"
OUT_MD = OUT_DIR / "Do_An_KWS_final_vi_2026_06_12.md"
OUT_DOCX = OUT_DIR / "Do_An_KWS_final_vi_2026_06_12.docx"

CAP620_CSV = ROOT / "results" / "cap620_16_pipeline_metrics_long.csv"
MICROSET_MD = ROOT / "reports" / "microset" / "result_table.md"
SERVER_FAR_MD = ROOT / "reports" / "server_far_metrics" / "server_far_metrics_summary.md"
CAP620_SUMMARY_MD = ROOT / "results" / "cap620_16_pipeline_test100_summary.md"

RUN_ID = "colab_mswc_cap620_flac_16pipe_e40_ep150_20260611_154517"

PAPER_ACC_1FAR = 82.0
PAPER_PARAMS = "128k"
PAPER_MACS = "29.4M"


@dataclass(frozen=True)
class TableBlock:
    caption: str
    headers: list[str]
    rows: list[list[str]]


@dataclass(frozen=True)
class FigureBlock:
    path: Path
    caption: str


def fnum(value: object, digits: int = 2) -> str:
    return f"{float(value):.{digits}f}"


def mean_pm(row: pd.Series, metric: str) -> str:
    std_key = f"{metric} std"
    if std_key in row and str(row[std_key]) != "":
        return f"{fnum(row[metric])} ± {fnum(row[std_key])}"
    return fnum(row[metric])


def pct(value: float | int | str, digits: int = 2) -> str:
    return f"{float(value):.{digits}f}%"


def label_model(value: str) -> str:
    return {
        "dscnn": "DSCNN-L",
        "edgespot_full": "EdgeSpotFull T4",
    }.get(value, value)


def label_frontend(value: str) -> str:
    value = str(value)
    if value.lower() in {"mel_pcen", "pcen"}:
        return "PCEN"
    if value.lower() == "mfcc":
        return "MFCC"
    return value


def label_loss(value: str) -> str:
    return {
        "triplet": "Triplet",
        "scaf": "SCAF",
        "ge2e": "GE2E",
        "scaf_ge2e": "SCAF+GE2E",
    }.get(value, value)


def pipeline_label(row: pd.Series) -> str:
    return f"{label_model(row['model_family'])} + {label_frontend(row['frontend'])} + {label_loss(row['loss'])}"


def load_cap620() -> pd.DataFrame:
    df = pd.read_csv(CAP620_CSV, encoding="utf-8-sig")
    numeric_cols = [
        "target_far_percent",
        "AUC",
        "AUC std",
        "EER",
        "EER std",
        "FRR@FAR",
        "FRR@FAR std",
        "Open-set ACC@FAR",
        "Open-set ACC@FAR std",
        "Keyword ACC",
        "Keyword ACC std",
        "Precision",
        "Precision std",
        "Recall",
        "Recall std",
        "F1",
        "F1 std",
    ]
    for col in numeric_cols:
        df[col] = pd.to_numeric(df[col], errors="coerce")
    df["pipeline"] = df.apply(pipeline_label, axis=1)
    df["model_label"] = df["model_family"].map(label_model)
    df["frontend_label"] = df["frontend"].map(label_frontend)
    df["loss_label"] = df["loss"].map(label_loss)
    return df


def row_lookup(df: pd.DataFrame, eval_name: str, model: str, frontend: str, loss: str) -> pd.Series:
    hit = df[
        (df["eval"] == eval_name)
        & (df["model_family"] == model)
        & (df["frontend_label"] == frontend)
        & (df["loss"] == loss)
    ]
    if hit.empty:
        raise KeyError((eval_name, model, frontend, loss))
    return hit.iloc[0]


def make_metric_table(df: pd.DataFrame, eval_name: str, title: str, sort_by: str | None = None) -> TableBlock:
    subset = df[df["eval"] == eval_name].copy()
    if sort_by:
        subset = subset.sort_values(sort_by, ascending=False)
    else:
        model_order = {"dscnn": 0, "edgespot_full": 1}
        frontend_order = {"MFCC": 0, "PCEN": 1}
        loss_order = {"Triplet": 0, "SCAF": 1, "GE2E": 2, "SCAF+GE2E": 3}
        subset["model_order"] = subset["model_family"].map(model_order)
        subset["frontend_order"] = subset["frontend_label"].map(frontend_order)
        subset["loss_order"] = subset["loss_label"].map(loss_order)
        subset = subset.sort_values(["model_order", "frontend_order", "loss_order"])

    rows = []
    for _, row in subset.iterrows():
        rows.append(
            [
                row["pipeline"],
                mean_pm(row, "Open-set ACC@FAR"),
                mean_pm(row, "AUC"),
                mean_pm(row, "EER"),
                mean_pm(row, "FRR@FAR"),
                mean_pm(row, "Keyword ACC"),
                mean_pm(row, "F1"),
            ]
        )
    return TableBlock(
        title,
        ["Pipeline", "ACC@FAR", "AUC", "EER", "FRR@FAR", "Keyword ACC", "F1"],
        rows,
    )


def make_top_table(df: pd.DataFrame, eval_name: str, top_n: int, caption: str) -> TableBlock:
    subset = df[df["eval"] == eval_name].sort_values("Open-set ACC@FAR", ascending=False).head(top_n)
    rows = []
    for rank, (_, row) in enumerate(subset.iterrows(), start=1):
        rows.append(
            [
                str(rank),
                row["pipeline"],
                mean_pm(row, "Open-set ACC@FAR"),
                mean_pm(row, "AUC"),
                mean_pm(row, "EER"),
                mean_pm(row, "F1"),
            ]
        )
    return TableBlock(caption, ["Rank", "Pipeline", "ACC@FAR", "AUC", "EER", "F1"], rows)


def make_delta_table(df: pd.DataFrame) -> TableBlock:
    test1 = df[df["eval"] == "test100_far1"]
    comparisons = [
        (
            "DSCNN-L, GE2E: PCEN so với MFCC",
            row_lookup(test1, "test100_far1", "dscnn", "MFCC", "ge2e"),
            row_lookup(test1, "test100_far1", "dscnn", "PCEN", "ge2e"),
        ),
        (
            "EdgeSpotFull T4, GE2E: PCEN so với MFCC",
            row_lookup(test1, "test100_far1", "edgespot_full", "MFCC", "ge2e"),
            row_lookup(test1, "test100_far1", "edgespot_full", "PCEN", "ge2e"),
        ),
        (
            "DSCNN-L, PCEN: GE2E so với Triplet",
            row_lookup(test1, "test100_far1", "dscnn", "PCEN", "triplet"),
            row_lookup(test1, "test100_far1", "dscnn", "PCEN", "ge2e"),
        ),
        (
            "EdgeSpotFull T4, PCEN: GE2E so với Triplet",
            row_lookup(test1, "test100_far1", "edgespot_full", "PCEN", "triplet"),
            row_lookup(test1, "test100_far1", "edgespot_full", "PCEN", "ge2e"),
        ),
        (
            "PCEN+GE2E: DSCNN-L so với EdgeSpotFull T4",
            row_lookup(test1, "test100_far1", "edgespot_full", "PCEN", "ge2e"),
            row_lookup(test1, "test100_far1", "dscnn", "PCEN", "ge2e"),
        ),
    ]
    rows = []
    for name, base, cand in comparisons:
        rows.append(
            [
                name,
                f"{float(cand['Open-set ACC@FAR']) - float(base['Open-set ACC@FAR']):+.2f}",
                f"{float(cand['AUC']) - float(base['AUC']):+.2f}",
                f"{float(cand['EER']) - float(base['EER']):+.2f}",
                f"{float(cand['F1']) - float(base['F1']):+.2f}",
            ]
        )
    return TableBlock(
        "Bảng delta trên GSC-test100@1%FAR. Giá trị dương ở ACC/AUC/F1 là tốt hơn; giá trị âm ở EER là tốt hơn.",
        ["So sánh", "ΔACC@1%FAR", "ΔAUC", "ΔEER", "ΔF1"],
        rows,
    )


def make_collapse_table(df: pd.DataFrame) -> TableBlock:
    subset = df[(df["eval"] == "test100_far1") & ((df["F1"] <= 1.0) | (df["FRR@FAR"] >= 99.0))]
    subset = subset.sort_values(["model_family", "frontend_label", "loss_label"])
    rows = []
    for _, row in subset.iterrows():
        rows.append(
            [
                row["pipeline"],
                fnum(row["AUC"]),
                fnum(row["EER"]),
                fnum(row["FRR@FAR"]),
                fnum(row["Open-set ACC@FAR"]),
                fnum(row["F1"]),
            ]
        )
    return TableBlock(
        "Các cấu hình có dấu hiệu collapse/reject-all trên GSC-test100@1%FAR.",
        ["Pipeline", "AUC", "EER", "FRR@1%FAR", "ACC@1%FAR", "F1"],
        rows,
    )


def make_paper_comparison_table(df: pd.DataFrame) -> TableBlock:
    test1 = df[df["eval"] == "test100_far1"]
    best_overall = test1.loc[test1["Open-set ACC@FAR"].idxmax()]
    best_edgespot = test1[test1["model_family"] == "edgespot_full"].sort_values("Open-set ACC@FAR", ascending=False).iloc[0]
    best_edgespot_auc = test1[test1["model_family"] == "edgespot_full"].sort_values("AUC", ascending=False).iloc[0]
    return TableBlock(
        "So sánh với mốc EdgeSpot-4 paper theo ACC@1%FAR.",
        ["Hệ thống", "Nguồn/profile", "Kích thước", "ACC@1%FAR", "Nhận xét"],
        [
            [
                "EdgeSpot-4 paper",
                "Paper EdgeSpot, 10-shot",
                f"{PAPER_PARAMS} params, {PAPER_MACS} MACs",
                pct(PAPER_ACC_1FAR, 1),
                "Mốc công bố trong paper; không phải kết quả chạy lại trong repo.",
            ],
            [
                best_overall["pipeline"],
                "Project, MSWC cap620 FLAC fixed",
                "~412.9k params",
                mean_pm(best_overall, "Open-set ACC@FAR"),
                "Nhỉnh hơn 82.0 rất nhẹ, nhưng model lớn hơn và sai số chuẩn chồng lấn.",
            ],
            [
                best_edgespot["pipeline"],
                "Project, MSWC cap620 FLAC fixed",
                "~130.6k params",
                mean_pm(best_edgespot, "Open-set ACC@FAR"),
                "Best compact EdgeSpot trong protocol cap620; chưa vượt paper.",
            ],
            [
                best_edgespot_auc["pipeline"],
                "Project, MSWC cap620 FLAC fixed",
                "~130.6k params",
                mean_pm(best_edgespot_auc, "Open-set ACC@FAR"),
                "AUC/EER/F1 tốt nhất trong nhóm EdgeSpot, nhưng ACC@1%FAR vẫn thấp hơn paper.",
            ],
        ],
    )


def md_table(table: TableBlock) -> str:
    lines = []
    lines.append(f"**{table.caption}**\n")
    lines.append("| " + " | ".join(table.headers) + " |")
    lines.append("| " + " | ".join(["---"] * len(table.headers)) + " |")
    lines.extend("| " + " | ".join(str(cell) for cell in row) + " |" for row in table.rows)
    return "\n".join(lines)


def section(title: str, level: int, paragraphs: Iterable[str]) -> str:
    marker = "#" * level
    parts = [f"{marker} {title}\n"]
    parts.extend(p.strip() + "\n" for p in paragraphs if p.strip())
    return "\n".join(parts)


def add_table_docx(document: Document, table: TableBlock) -> None:
    caption = document.add_paragraph(table.caption)
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in caption.runs:
        run.italic = True

    tbl = document.add_table(rows=1, cols=len(table.headers))
    tbl.style = "Table Grid"
    tbl.autofit = True
    for idx, header in enumerate(table.headers):
        cell = tbl.rows[0].cells[idx]
        cell.text = header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
    for row in table.rows:
        cells = tbl.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = str(value)
    document.add_paragraph("")


def add_figure_docx(document: Document, fig: FigureBlock) -> None:
    if not fig.path.exists():
        return
    document.add_picture(str(fig.path), width=Inches(6.0))
    caption = document.add_paragraph(fig.caption)
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in caption.runs:
        run.italic = True
    document.add_paragraph("")


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_default_styles(document: Document) -> None:
    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(13)
    normal.paragraph_format.line_spacing = 1.35
    normal.paragraph_format.space_after = Pt(6)

    for style_name in ["Heading 1", "Heading 2", "Heading 3"]:
        style = styles[style_name]
        style.font.name = "Times New Roman"
        style.font.bold = True
        if style_name == "Heading 1":
            style.font.size = Pt(16)
        elif style_name == "Heading 2":
            style.font.size = Pt(14)
        else:
            style.font.size = Pt(13)


def set_page_layout(document: Document) -> None:
    for sec in document.sections:
        sec.top_margin = Cm(2.5)
        sec.bottom_margin = Cm(2.5)
        sec.left_margin = Cm(3.0)
        sec.right_margin = Cm(2.0)


def add_field(paragraph, instruction: str) -> None:
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")

    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = instruction

    fld_separate = OxmlElement("w:fldChar")
    fld_separate.set(qn("w:fldCharType"), "separate")

    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")

    run._r.append(fld_begin)
    run._r.append(instr_text)
    run._r.append(fld_separate)
    run._r.append(fld_end)


def add_page_numbers(document: Document) -> None:
    for sec in document.sections:
        footer = sec.footer
        paragraph = footer.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.add_run("Trang ")
        add_field(paragraph, "PAGE")


def add_cover(document: Document) -> None:
    for text, size, bold in [
        ("TRƯỜNG/VIỆN: [ĐIỀN THEO MẪU]", 13, True),
        ("KHOA/CENTRE: [ĐIỀN THEO MẪU]", 13, True),
    ]:
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text)
        r.bold = bold
        r.font.size = Pt(size)

    document.add_paragraph("")
    document.add_paragraph("")
    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("FEW-SHOT OPEN-SET KEYWORD SPOTTING\nDỰA TRÊN EMBEDDING VÀ PROTOTYPE INFERENCE")
    r.bold = True
    r.font.size = Pt(18)
    r.font.color.rgb = RGBColor(0, 0, 0)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = subtitle.add_run("BẢN THESIS TIẾNG VIỆT - DRAFT 2026-06-12")
    r.bold = True
    r.font.size = Pt(14)

    document.add_paragraph("")
    info = [
        "Sinh viên: [Điền tên sinh viên]",
        "Mã sinh viên: [Điền mã sinh viên]",
        "Ngành/chương trình: [Điền ngành]",
        "Giảng viên hướng dẫn: [Điền tên giảng viên hướng dẫn]",
        "Nguồn số liệu chính: MSWC cap620 FLAC fixed 16-pipeline, GSC-test100",
    ]
    for line in info:
        p = document.add_paragraph(line)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    document.add_paragraph("")
    p = document.add_paragraph("Hà Nội, 2026")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_page_break()


def add_toc_placeholder(document: Document) -> None:
    document.add_heading("Mục lục", level=1)
    p = document.add_paragraph()
    add_field(p, 'TOC \\o "1-3" \\h \\z \\u')
    document.add_page_break()


def add_section_docx(document: Document, title: str, level: int, paragraphs: Iterable[str]) -> None:
    document.add_heading(title, level=level)
    for para in paragraphs:
        if not para.strip():
            continue
        p = document.add_paragraph(para.strip())
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def read_text_if_exists(path: Path) -> str:
    return path.read_text(encoding="utf-8") if path.exists() else ""


def make_figures(df: pd.DataFrame) -> list[FigureBlock]:
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    plt.rcParams.update({"font.size": 9, "figure.dpi": 140})

    test1 = df[df["eval"] == "test100_far1"].copy()
    top = test1.sort_values("Open-set ACC@FAR", ascending=False).head(8).iloc[::-1]
    fig, ax = plt.subplots(figsize=(8.4, 4.8))
    ax.barh(top["pipeline"], top["Open-set ACC@FAR"], color="#2f6f73")
    ax.axvline(PAPER_ACC_1FAR, color="#c33d2e", linestyle="--", linewidth=1.5, label="EdgeSpot-4 paper 82.0")
    ax.set_xlabel("ACC@1%FAR (%)")
    ax.set_title("Top pipelines on GSC-test100 @ 1% FAR")
    ax.set_xlim(65, 88)
    ax.legend(loc="lower right")
    fig.tight_layout()
    rank_path = ASSET_DIR / "cap620_top8_acc1far.png"
    fig.savefig(rank_path, bbox_inches="tight")
    plt.close(fig)

    heat = test1.pivot_table(
        index=["model_label", "frontend_label"],
        columns="loss_label",
        values="Open-set ACC@FAR",
        aggfunc="mean",
    )
    heat = heat[["Triplet", "SCAF", "GE2E", "SCAF+GE2E"]]
    fig, ax = plt.subplots(figsize=(7.2, 3.6))
    im = ax.imshow(heat.values, cmap="viridis", vmin=45, vmax=85)
    ax.set_xticks(range(len(heat.columns)), labels=heat.columns)
    ax.set_yticks(range(len(heat.index)), labels=[f"{a} / {b}" for a, b in heat.index])
    for i in range(heat.shape[0]):
        for j in range(heat.shape[1]):
            ax.text(j, i, f"{heat.values[i, j]:.1f}", ha="center", va="center", color="white", fontsize=8)
    ax.set_title("ACC@1%FAR interaction: model, frontend and loss")
    fig.colorbar(im, ax=ax, label="ACC@1%FAR (%)")
    fig.tight_layout()
    heat_path = ASSET_DIR / "cap620_acc1far_heatmap.png"
    fig.savefig(heat_path, bbox_inches="tight")
    plt.close(fig)

    test5 = df[df["eval"] == "test100_far5"]
    best_dscnn = test1.loc[test1[test1["model_family"] == "dscnn"]["Open-set ACC@FAR"].idxmax()]
    best_edge = test1.loc[test1[test1["model_family"] == "edgespot_full"]["Open-set ACC@FAR"].idxmax()]
    fig, ax = plt.subplots(figsize=(7.0, 4.0))
    labels = ["EdgeSpot-4\npaper", "Best overall\n(project)", "Best EdgeSpot\n(project)"]
    vals = [PAPER_ACC_1FAR, best_dscnn["Open-set ACC@FAR"], best_edge["Open-set ACC@FAR"]]
    colors = ["#c33d2e", "#2f6f73", "#7266ba"]
    bars = ax.bar(labels, vals, color=colors)
    ax.set_ylabel("ACC@1%FAR (%)")
    ax.set_ylim(76, 84.5)
    ax.set_title("Comparison boundary with EdgeSpot-4 paper")
    for bar, val in zip(bars, vals):
        ax.text(bar.get_x() + bar.get_width() / 2, val + 0.08, f"{val:.2f}", ha="center", va="bottom")
    fig.tight_layout()
    cmp_path = ASSET_DIR / "edgespot4_comparison_acc1far.png"
    fig.savefig(cmp_path, bbox_inches="tight")
    plt.close(fig)

    return [
        FigureBlock(rank_path, "Hình 1. Top 8 pipeline theo ACC@1%FAR trên GSC-test100; đường đứt là mốc EdgeSpot-4 paper."),
        FigureBlock(heat_path, "Hình 2. Tương tác giữa backbone, frontend và loss trong thí nghiệm cap620 fixed."),
        FigureBlock(cmp_path, "Hình 3. Ranh giới so sánh với EdgeSpot-4 paper: best overall khác với best compact EdgeSpot."),
    ]


def build_content(df: pd.DataFrame) -> tuple[str, list[TableBlock], list[FigureBlock], list[tuple[str, int, list[str], list[TableBlock], list[FigureBlock]]]]:
    total_rows = len(df)
    unique_pipelines = df["run_tag"].nunique()
    all_ok = bool(
        (df[["train_status", "dev30_far1_status", "test100_far1_status", "test100_far5_status"]] == "ok")
        .all()
        .all()
    )
    all_ok_text = "ok" if all_ok else "chưa ok"
    test1 = df[df["eval"] == "test100_far1"].copy()
    test5 = df[df["eval"] == "test100_far5"].copy()
    best1 = test1.loc[test1["Open-set ACC@FAR"].idxmax()]
    best5 = test5.loc[test5["Open-set ACC@FAR"].idxmax()]
    edge_best1 = test1[test1["model_family"] == "edgespot_full"].sort_values("Open-set ACC@FAR", ascending=False).iloc[0]
    edge_best_auc = test1[test1["model_family"] == "edgespot_full"].sort_values("AUC", ascending=False).iloc[0]
    dscnn_best1 = test1[test1["model_family"] == "dscnn"].sort_values("Open-set ACC@FAR", ascending=False).iloc[0]

    summary_tables = [
        make_top_table(df, "test100_far1", 6, "Top pipeline trên GSC-test100@1%FAR."),
        make_top_table(df, "test100_far5", 6, "Top pipeline trên GSC-test100@5%FAR."),
        make_delta_table(df),
        make_paper_comparison_table(df),
        TableBlock(
            "Các mốc thực nghiệm phụ để đặt kết quả cap620 vào bối cảnh. Các dòng này không dùng để ranking chung vì khác protocol/data profile.",
            ["Profile", "Pipeline", "Split", "ACC@1%FAR", "ACC@5%FAR", "AUC", "EER", "F1", "Vai trò"],
            [
                [
                    "Microset",
                    "EdgeSpotFull T4 + PCEN + SCAF+GE2E",
                    "GSC-test100",
                    "84.61%",
                    "86.12%",
                    "95.61%",
                    "11.54%",
                    "82.41%",
                    "Evidence ban đầu cho hướng compact hybrid trên dataset nhỏ.",
                ],
                [
                    "Top500 epoch13",
                    "EdgeSpotFull T4 + PCEN + SCAF+GE2E",
                    "GSC-test100",
                    "85.62 ± 1.04",
                    "88.79 ± 0.66",
                    "95.34 ± 0.40",
                    "11.51 ± 0.76",
                    "82.45 ± 1.08",
                    "Artifact riêng cho thấy tiềm năng EdgeSpot+SCAF+GE2E, không thay cap620 fixed.",
                ],
                [
                    "cap620 fixed",
                    "DSCNN-L + PCEN + GE2E",
                    "GSC-test100",
                    "82.34 ± 1.19",
                    "86.57 ± 0.75",
                    "92.42 ± 0.54",
                    "14.89 ± 0.84",
                    "77.75 ± 1.15",
                    "Evidence chính của thesis cho ablation 16 pipeline.",
                ],
                [
                    "cap620 fixed",
                    "EdgeSpotFull T4 + PCEN + GE2E",
                    "GSC-test100",
                    "79.98 ± 0.98",
                    "83.16 ± 0.82",
                    "87.23 ± 0.75",
                    "20.23 ± 0.96",
                    "70.68 ± 1.23",
                    "Best compact EdgeSpot theo ACC@1%FAR trong cap620 fixed.",
                ],
            ],
        ),
    ]
    full_tables = [
        make_metric_table(df, "test100_far1", "Toàn bộ 16 pipeline trên GSC-test100@1%FAR."),
        make_metric_table(df, "test100_far5", "Toàn bộ 16 pipeline trên GSC-test100@5%FAR."),
        make_metric_table(df, "dev30_far1", "Kết quả GSC-dev30@1%FAR sau huấn luyện."),
        make_collapse_table(df),
    ]
    figures = make_figures(df)

    abstract = [
        "Đồ án này nghiên cứu bài toán few-shot open-set keyword spotting (KWS), trong đó hệ thống phải nhận diện các từ khóa mới từ một số ít mẫu enrollment và đồng thời từ chối các âm thanh không thuộc tập từ khóa đã đăng ký. Khác với KWS closed-set, bài toán open-set yêu cầu kiểm soát false acceptance rate (FAR) vì một false accept có thể kích hoạt sai lệnh thoại trong hệ thống thực tế.",
        f"Đóng góp thực nghiệm chính của bản thesis là một thí nghiệm fixed 16-pipeline trên MSWC English cap620 FLAC với cùng dữ liệu, cùng lịch huấn luyện và cùng protocol đánh giá. Run `{RUN_ID}` tạo {total_rows} dòng metric cho {unique_pipelines} pipeline, gồm `dev30_far1`, `test100_far1` và `test100_far5`; trạng thái hoàn tất của tất cả stage là `{all_ok_text}`.",
        f"Kết quả chính cho thấy `{best1['pipeline']}` là cấu hình accuracy-oriented tốt nhất, đạt ACC@1%FAR = {mean_pm(best1, 'Open-set ACC@FAR')}, AUC = {mean_pm(best1, 'AUC')}, EER = {mean_pm(best1, 'EER')} và F1 = {mean_pm(best1, 'F1')} trên GSC-test100. Ở FAR=5%, cùng cấu hình đạt ACC@5%FAR = {mean_pm(best5, 'Open-set ACC@FAR')}. Trong nhóm compact EdgeSpotFull T4, `{edge_best1['pipeline']}` đạt ACC@1%FAR cao nhất ({mean_pm(edge_best1, 'Open-set ACC@FAR')}), trong khi `{edge_best_auc['pipeline']}` tốt hơn về AUC/EER/F1.",
        "So với mốc EdgeSpot-4 paper đầu năm 2026, kết quả tốt nhất của project xấp xỉ và nhỉnh nhẹ 82.0% ACC@1%FAR, nhưng đó là DSCNN-L lớn hơn. Bản EdgeSpotFull T4 compact trong protocol cap620 hiện chưa vượt paper. Vì vậy, claim đúng là project đã đạt mức cạnh tranh với paper bằng cấu hình accuracy-oriented, còn hướng compact cần thêm tối ưu như distillation, objective phù hợp hơn hoặc tuning loss.",
    ]

    sections: list[tuple[str, int, list[str], list[TableBlock], list[FigureBlock]]] = [
        (
            "Lời cảm ơn",
            1,
            [
                "Em xin gửi lời cảm ơn tới giảng viên hướng dẫn, các thầy cô và các anh chị đã hỗ trợ về chuyên môn, tài nguyên tính toán và góp ý trong quá trình thực hiện đồ án. Các thí nghiệm trong đồ án yêu cầu nhiều lần chạy trên Colab và server, vì vậy sự hỗ trợ về môi trường chạy và phản hồi kỹ thuật có vai trò quan trọng để hoàn thiện kết quả.",
                "Em cũng xin cảm ơn gia đình và bạn bè đã động viên trong quá trình làm việc. Bản thảo này được viết theo hướng có thể tiếp tục chỉnh sửa theo mẫu hình thức của trường, trong đó các phần bìa, thông tin sinh viên và format citation cần được cập nhật theo yêu cầu chính thức trước khi nộp.",
            ],
            [],
            [],
        ),
        ("Tóm tắt", 1, abstract, [], []),
        (
            "Danh mục thuật ngữ viết tắt",
            1,
            [
                "KWS: Keyword Spotting. FAR: False Acceptance Rate. FRR: False Rejection Rate. EER: Equal Error Rate. AUC: Area Under the ROC Curve. GSC: Google Speech Commands. MSWC: Multilingual Spoken Words Corpus. PCEN: Per-Channel Energy Normalization. MFCC: Mel-Frequency Cepstral Coefficients. GE2E: Generalized End-to-End. SCAF: Sub-center ArcFace.",
            ],
            [
                TableBlock(
                    "Các thuật ngữ chính dùng trong báo cáo.",
                    ["Thuật ngữ", "Ý nghĩa trong đồ án"],
                    [
                        ["Few-shot KWS", "Nhận diện từ khóa mới từ số ít mẫu support/enrollment."],
                        ["Open-set rejection", "Từ chối query không thuộc các keyword đã enroll."],
                        ["Prototype", "Vector đại diện của keyword, tính bằng trung bình embedding support samples."],
                        ["ACC@1%FAR", "Open-set accuracy tại ngưỡng vận hành giới hạn false accept ở 1%."],
                        ["Test100", "Đánh giá trung bình qua 100 repeated few-shot episodes."],
                    ],
                )
            ],
            [],
        ),
        (
            "Chương 1. Giới thiệu",
            1,
            [
                "Keyword Spotting là bài toán phát hiện một hoặc nhiều từ khóa trong tín hiệu âm thanh ngắn. Trong các hệ thống trợ lý giọng nói, thiết bị thông minh hoặc điều khiển rảnh tay, KWS thường là thành phần đầu vào quyết định khi nào hệ thống cần phản hồi. Một hệ thống KWS thực tế không chỉ cần nhận diện đúng keyword mà còn cần tránh kích hoạt sai khi người dùng nói từ khác hoặc khi môi trường có nhiễu.",
                "Nhiều hệ thống KWS truyền thống được thiết kế như bài toán closed-set classification: mô hình chọn một nhãn trong tập keyword cố định. Cách này hiệu quả khi tập từ khóa không đổi, nhưng kém linh hoạt khi người dùng muốn thêm từ khóa cá nhân hóa chỉ bằng vài mẫu. Few-shot KWS giải quyết vấn đề bằng cách học embedding space, sau đó thêm keyword mới bằng prototype thay vì huấn luyện lại toàn bộ classifier.",
                "Thách thức chính của few-shot KWS nằm ở open-set setting. Query audio có thể là keyword đã enroll, một từ gần âm, một từ ngoài vocabulary hoặc silence/noise. Nếu chỉ dùng nearest prototype mà không có cơ chế threshold, hệ thống sẽ luôn ép query vào một keyword, dẫn đến false accept. Vì vậy, đồ án tập trung vào pipeline embedding + prototype + threshold, trong đó metric chính là ACC tại các FAR cố định.",
                "Mục tiêu nghiên cứu là trả lời bốn câu hỏi: (1) frontend nào phù hợp cho few-shot open-set KWS, MFCC hay PCEN; (2) loss nào phù hợp với prototype inference, Triplet, SCAF, GE2E hay SCAF+GE2E; (3) backbone nào tốt hơn giữa DSCNN-L và EdgeSpotFull T4; (4) kết quả của project so với EdgeSpot-4 paper nên được claim như thế nào cho đúng.",
                "Đóng góp của đồ án gồm: xây dựng pipeline few-shot open-set KWS end-to-end; triển khai hai nhóm backbone DSCNN-L và EdgeSpotFull T4; đánh giá có hệ thống 16 pipeline trên cùng protocol cap620 fixed; phân tích vì sao PCEN/GE2E tốt, vì sao SCAF collapse trong profile lớn; và tạo demo web phục vụ enrollment, single detection, long-audio analysis và open-set calibration.",
            ],
            [],
            [],
        ),
        (
            "Chương 2. Cơ sở lý thuyết và công trình liên quan",
            1,
            [
                "Trong embedding-based KWS, encoder biến audio thành một vector có chiều thấp hơn. Các mẫu cùng keyword được kỳ vọng nằm gần nhau trong embedding space, còn các keyword khác nhau nằm xa nhau. Khi người dùng enroll một keyword, hệ thống lấy k mẫu support, chạy qua encoder và tính trung bình embedding để tạo prototype. Query được so sánh với các prototype bằng distance hoặc similarity.",
                "MFCC là frontend cổ điển trong speech processing. Nó nén phổ mel thành cepstral coefficients, giúp giảm chiều và tạo representation gọn. Tuy nhiên, MFCC có thể làm mất một phần chi tiết time-frequency và nhạy với điều kiện thu âm. Trong đồ án, MFCC được giữ như baseline và ablation để kiểm tra giá trị của frontend truyền thống.",
                "PCEN là frontend theo hướng chuẩn hóa năng lượng theo kênh, có khả năng giảm ảnh hưởng của biến thiên âm lượng và nhiễu nền. Với few-shot KWS, support và query có thể đến từ speaker hoặc thiết bị thu khác nhau, nên frontend ổn định về năng lượng có ý nghĩa trực tiếp với distance trong embedding space. Kết quả cap620 cho thấy PCEN là thành phần có ảnh hưởng dương rõ nhất.",
                "Triplet loss tối ưu quan hệ tương đối giữa anchor, positive và negative. Loss này phù hợp với metric learning vì nó trực tiếp đẩy mẫu cùng lớp lại gần và mẫu khác lớp ra xa. Tuy nhiên, hiệu quả phụ thuộc vào mining strategy; nếu negative quá dễ, gradient yếu, còn nếu quá khó có thể gây training không ổn định.",
                "GE2E loss dùng centroid/prototype trong chính objective huấn luyện. Trong mỗi episode, một phần mẫu của mỗi class tạo centroid, phần query còn lại được phân loại theo similarity với các centroid. Cơ chế này gần với inference thật của few-shot KWS, nên GE2E thường phù hợp với prototype inference hơn loss phân loại thuần túy.",
                "SCAF là biến thể Sub-center ArcFace. Mỗi class có nhiều sub-center để hấp thụ nhiễu nội lớp, đồng thời dùng angular margin để tăng phân tách. Ý tưởng này hấp dẫn với dữ liệu speech có nhiều speaker, nhưng khi số lớp train lên tới hàng chục nghìn, classifier head và scale/margin có thể trở thành nguồn bất ổn nếu không tune kỹ.",
                "EdgeSpot là hướng mô hình nhỏ gọn cho few-shot KWS. Trong đồ án, EdgeSpotFull T4 được triển khai như compact candidate với khoảng 130.6k tham số, nhỏ hơn DSCNN-L khoảng ba lần. Điểm cần nhấn mạnh là project không claim reproduction đầy đủ của paper EdgeSpot; project dùng EdgeSpot-style backbone trong protocol riêng và so sánh claim ở mức metric công bố.",
            ],
            [],
            [],
        ),
        (
            "Chương 3. Thiết kế hệ thống và phương pháp",
            1,
            [
                "Pipeline tổng quát gồm sáu bước: chuẩn hóa audio về mono 16 kHz và độ dài khoảng 1 giây; trích xuất MFCC hoặc mel-PCEN; chạy encoder để lấy embedding; L2-normalize embedding; tạo prototype từ support samples; và đưa ra quyết định accept/reject bằng ngưỡng score tại target FAR.",
                "DSCNN-L được cài đặt bằng depthwise separable convolution. Theo code `src/models/dscnn.py`, model L dùng 276 channels, một convolution ban đầu và 5 depthwise-separable blocks. Input mặc định cho MFCC là `(47, 10)`, còn các thí nghiệm mel/PCEN dùng dạng time-frequency map lớn hơn `(40, 101)`. Embedding đầu ra của DSCNN-L có 276 chiều.",
                "EdgeSpotFull T4 được cài đặt trong `src/models/edgespot_full.py`. Model dùng trainable PCEN, stem convolution, các fused temporal/BC-ResNet-style blocks, depthwise temporal positional convolution, single-head attention và head tạo embedding 64 chiều. MFCC vẫn được hỗ trợ như ablation, nhưng đường thiết kế tự nhiên của EdgeSpotFull T4 là mel/PCEN.",
                "Training dùng episodic sampling. Mỗi episode lấy 30 class và 10 sample mỗi class trong run cap620 fixed. Với 150 episode/epoch và 40 epoch, số sample occurrence theo episode là khoảng 1.8 triệu. Đây không phải epoch supervised quét hết 2.99 triệu file; do đó kích thước manifest lớn không đồng nghĩa toàn bộ file đều được quan sát đều như nhau.",
                "Checkpoint tốt nhất không chọn theo train loss mà chọn theo GSC-dev ACC@1%FAR. Trong run cap620, cứ mỗi 5 epoch mô hình được evaluate 3 runs trên GSC-dev với k=10. Cách chọn này bám sát mục tiêu open-set hơn vì train loss thấp chưa chắc tạo threshold tốt ở FAR thấp.",
                "Evaluation chính dùng protocol `gsc_edgespot_exact`. Tập target gồm 10 command words của GSC cộng với silence thật; negative là 25 spoken words còn lại ngoài 10 command target. Mỗi run lấy 10 support samples mỗi keyword để tạo prototype, sau đó đánh giá query samples. Final dev dùng 30 runs, final test dùng 100 runs.",
                "Các metric chính gồm AUC, EER, FRR@FAR, ACC@FAR, Keyword ACC, Precision, Recall và F1. ACC@FAR là open-set multiclass accuracy tại threshold sao cho FAR không vượt target. FRR@FAR cho biết tỷ lệ positive keyword bị reject ở ngưỡng đó. Vì vậy, một pipeline có ACC nhìn cao nhưng FRR=100% và F1=0 không thể xem là tốt.",
            ],
            [
                TableBlock(
                    "Cấu hình training cố định của thí nghiệm cap620.",
                    ["Trường", "Giá trị"],
                    [
                        ["Data profile", "MSWC English cap620 FLAC"],
                        ["Train files", "2,989,780"],
                        ["Validation files", "52,399"],
                        ["Train words", "37,387"],
                        ["Validation words", "763"],
                        ["Epochs", "40"],
                        ["Episodes/epoch", "150"],
                        ["Episode shape", "30 classes × 10 samples"],
                        ["Optimizer", "Adam, lr=0.001, weight_decay=0.0001"],
                        ["Scheduler", "CosineAnnealingWarmRestarts"],
                        ["Checkpoint selection", "GSC-dev ACC@1%FAR, every 5 epochs, 3 runs"],
                        ["Final evaluation", "dev30@1%FAR, test100@1%FAR, test100@5%FAR"],
                    ],
                ),
                TableBlock(
                    "Ma trận 16 pipeline trong thí nghiệm fixed.",
                    ["Backbone", "Frontend", "Loss"],
                    [
                        ["DSCNN-L", "MFCC", "Triplet / SCAF / GE2E / SCAF+GE2E"],
                        ["DSCNN-L", "PCEN", "Triplet / SCAF / GE2E / SCAF+GE2E"],
                        ["EdgeSpotFull T4", "MFCC", "Triplet / SCAF / GE2E / SCAF+GE2E"],
                        ["EdgeSpotFull T4", "PCEN", "Triplet / SCAF / GE2E / SCAF+GE2E"],
                    ],
                ),
            ],
            [],
        ),
        (
            "Chương 4. Thực nghiệm",
            1,
            [
                "Thí nghiệm chính của thesis là fixed 16-pipeline cap620 FLAC. Tất cả pipeline dùng cùng dữ liệu, cùng lịch huấn luyện và cùng protocol evaluation. Đây là điểm khác với các run lịch sử như Microset, Top500 hoặc manifest20/manifest50, vốn có giá trị bối cảnh nhưng không nên trộn vào cùng một ranking final.",
                "Dữ liệu train là MSWC English với giới hạn tối đa 620 clip mỗi từ. Audio được tải ở OPUS, chuyển sang FLAC và xóa OPUS để giảm áp lực disk trong Colab. Artifact quan trọng được sync lên Drive gồm checkpoint, results, reports, logs, configs và splits; audio clips không sync lên Drive.",
                "GSC v2 chỉ dùng để evaluate, không dùng làm tập train chính cho các mô hình cap620. Điều này tạo một cross-dataset setting: encoder học từ MSWC nhưng được kiểm tra trên GSC command words. Nếu kết quả tốt, điều đó cho thấy embedding có khả năng chuyển giao sang tập command khác.",
                "Tất cả 16 pipeline hoàn tất train và evaluate. File evidence chính là `results/cap620_16_pipeline_metrics_long.csv`; file này có 48 dòng metric tương ứng 16 pipeline × 3 eval settings. Các cột status đều là `ok`, do đó không có pipeline bị thiếu final test trong bảng chính.",
                "Bên cạnh thí nghiệm cap620, thesis vẫn nhắc Microset và Top500 như các mốc phát triển. Microset cho thấy SCAF+GE2E từng có tín hiệu tốt trên setting nhỏ. Top500 epoch13 là artifact EdgeSpotFull T4 + PCEN + SCAF+GE2E có ACC@1%FAR cao, nhưng nó thuộc profile khác và không thay thế kết luận cap620 fixed.",
            ],
            [summary_tables[4], *summary_tables[:2]],
            figures[:2],
        ),
        (
            "Chương 5. Kết quả và thảo luận",
            1,
            [
                f"Cấu hình tốt nhất toàn bộ là `{best1['pipeline']}`. Trên GSC-test100@1%FAR, cấu hình này đạt ACC@1%FAR = {mean_pm(best1, 'Open-set ACC@FAR')}, AUC = {mean_pm(best1, 'AUC')}, EER = {mean_pm(best1, 'EER')}, FRR@1%FAR = {mean_pm(best1, 'FRR@FAR')}, Keyword ACC = {mean_pm(best1, 'Keyword ACC')} và F1 = {mean_pm(best1, 'F1')}. Ở FAR=5%, nó đạt ACC@5%FAR = {mean_pm(best5, 'Open-set ACC@FAR')}.",
                "PCEN là frontend ổn định nhất trong thí nghiệm. Với GE2E, đổi MFCC sang PCEN tăng ACC@1%FAR của DSCNN-L thêm 5.26 điểm và tăng F1 thêm 9.25 điểm. Với EdgeSpotFull T4, mức tăng còn lớn hơn: ACC tăng 9.22 điểm và F1 tăng 21.86 điểm. Điều này cho thấy EdgeSpot-style backbone đặc biệt phụ thuộc vào mel/PCEN map thay vì MFCC nén cepstral.",
                "GE2E phù hợp nhất với DSCNN-L vì objective centroid/prototype khớp với inference. Trên DSCNN-L + PCEN, GE2E vượt Triplet 2.36 điểm ACC@1%FAR, 1.77 điểm AUC và 3.61 điểm F1. Với capacity lớn hơn, DSCNN-L tận dụng tốt GE2E để hình thành embedding space có cấu trúc centroid rõ.",
                "Trong nhóm EdgeSpotFull T4, kết luận tinh hơn. PCEN + GE2E nhỉnh PCEN + Triplet 0.40 điểm ACC@1%FAR, nhưng PCEN + Triplet lại tốt hơn ở AUC, EER và F1. Điều này cho thấy nếu mục tiêu là compact model có calibration linh hoạt, Triplet vẫn rất đáng giữ lại, thay vì chỉ chọn theo một operating point.",
                "SCAF và SCAF+GE2E collapse ở nhiều pipeline cap620. Dấu hiệu gồm AUC khoảng 50%, EER khoảng 50%, FRR@FAR = 100% và F1 = 0. Trong trường hợp này, ACC khoảng 69.44% không có ý nghĩa tốt vì model gần như reject toàn bộ positive queries nhưng vẫn đúng trên nhiều unknown samples. Đây là ví dụ rõ ràng vì sao open-set thesis phải báo cáo FRR và F1, không chỉ ACC.",
                "Nguyên nhân hợp lý của SCAF collapse là mismatch giữa angular classification head và setting 37k train words. SCAF cần classifier head với rất nhiều class và sub-centers; trong khi episodic batch chỉ quan sát 30 class mỗi episode. Nếu scale/margin/loss weight không phù hợp, gradient classification có thể dominating và phá vỡ cấu trúc embedding prototype.",
                "So sánh backbone cho thấy DSCNN-L tốt hơn về accuracy tuyệt đối, còn EdgeSpotFull T4 có lợi thế compact. Best DSCNN đạt 82.34% ACC@1%FAR, best EdgeSpot đạt 79.98%, chênh 2.36 điểm. Với tham số khoảng 412.9k so với 130.6k, lựa chọn cuối phụ thuộc vào mục tiêu: accuracy hay edge deployment.",
            ],
            [summary_tables[2], full_tables[3], full_tables[0], full_tables[1]],
            [figures[2]],
        ),
        (
            "Chương 6. So sánh với EdgeSpot-4 paper",
            1,
            [
                f"Paper EdgeSpot đầu năm 2026 báo cáo EdgeSpot-4 đạt 10-shot ACC@1%FAR = {PAPER_ACC_1FAR:.1f}% với {PAPER_PARAMS} tham số và {PAPER_MACS} MACs. Đây là mốc quan trọng vì metric cùng là ACC@1%FAR trong few-shot KWS. Tuy nhiên, project không chạy reproduction đầy đủ cùng code, split và recipe của paper, nên phần so sánh phải được viết như benchmark boundary thay vì claim tái lập paper.",
                f"Best overall của project trong cap620 fixed là `{dscnn_best1['pipeline']}`, đạt {mean_pm(dscnn_best1, 'Open-set ACC@FAR')} ACC@1%FAR. Con số này nhỉnh hơn 82.0 của paper rất nhẹ, nhưng model lớn hơn khoảng ba lần và khoảng tin cậy chồng lấn. Do đó, câu claim chuẩn là project đạt mức cạnh tranh/xấp xỉ EdgeSpot-4 bằng cấu hình accuracy-oriented.",
                f"Trong nhóm compact EdgeSpotFull T4 của project, best theo ACC@1%FAR là `{edge_best1['pipeline']}` với {mean_pm(edge_best1, 'Open-set ACC@FAR')}. Kết quả này thấp hơn mốc paper khoảng {PAPER_ACC_1FAR - float(edge_best1['Open-set ACC@FAR']):.2f} điểm. Vì vậy không nên viết rằng EdgeSpotFull T4 cap620 đã vượt EdgeSpot-4 paper.",
                "Top500 epoch13 của project có artifact EdgeSpotFull T4 + PCEN + SCAF+GE2E đạt 85.62% ACC@1%FAR, cao hơn 82.0. Tuy nhiên, đây là profile huấn luyện khác, không thuộc fixed cap620 16-pipeline. Có thể đưa vào như evidence riêng cho tiềm năng của hướng EdgeSpot+SCAF+GE2E, nhưng không dùng để thay thế kết luận cap620.",
                "Hướng hợp lý để compact EdgeSpotFull T4 vượt paper là bổ sung distillation hoặc teacher-guided objective, tune Triplet/GE2E cho EdgeSpot, và chỉ quay lại SCAF sau khi có ablation về margin, scale, loss weight và warmup. Chạy thêm dữ liệu mà không sửa objective có thể không giải quyết được collapse.",
            ],
            [summary_tables[3]],
            [],
        ),
        (
            "Chương 7. Demo system và triển khai",
            1,
            [
                "Demo web của project minh họa pipeline few-shot open-set KWS ngoài các bảng metric. Người dùng có thể enroll keyword bằng audio mẫu, chạy single detection, phân tích long audio và thử open-set rejection. Demo hiển thị top candidates, distance, threshold, margin và lý do accept/reject để hỗ trợ phân tích lỗi.",
                "Backend demo tải checkpoint, chọn frontend phù hợp với metadata checkpoint, trích xuất feature và xây dựng prototype từ enrollment cache. Khi người dùng đổi model profile, hệ thống cần rebuild hoặc clear enrollment vì embedding space của mỗi model khác nhau. Đây là điểm quan trọng để tránh dùng prototype của model cũ cho model mới.",
                "Open-set UI sampled evaluation chỉ có giá trị demo/debug. Kết quả nghiên cứu trong thesis phải dựa trên `gsc_edgespot_exact` dev/test với số runs rõ ràng. Vì vậy, nếu UI cho kết quả tốt nhưng test100 không xác nhận, không được dùng UI để claim final performance.",
                "Long-audio flow giúp kiểm tra các lỗi thực tế như miss do threshold, nhầm từ gần âm, VAD/cooldown skip hoặc lệch timing. Các kết quả này có giá trị engineering và demo, nhưng một benchmark streaming chính thức cần thêm latency, false alarm per hour và miss rate trên audio liên tục.",
            ],
            [],
            [],
        ),
        (
            "Chương 8. Kết luận và hướng phát triển",
            1,
            [
                "Đồ án đã xây dựng và đánh giá một pipeline few-shot open-set keyword spotting dựa trên embedding và prototype inference. Hệ thống có khả năng thêm keyword mới bằng số ít support samples, sau đó nhận diện hoặc reject query audio bằng distance threshold tại target FAR.",
                f"Thí nghiệm fixed 16-pipeline cap620 là evidence mạnh nhất hiện tại. Kết quả kết luận rằng PCEN là frontend nên dùng mặc định, GE2E là loss tốt nhất cho DSCNN-L, Triplet/GE2E là hai lựa chọn đáng giữ cho EdgeSpotFull T4, và SCAF/SCAF+GE2E cần tuning lại trước khi dùng trên profile 37k words.",
                "Về claim với paper, project đạt mức cạnh tranh với EdgeSpot-4 bằng DSCNN-L + PCEN + GE2E, nhưng compact EdgeSpotFull T4 cap620 chưa vượt EdgeSpot-4. Đây là ranh giới claim quan trọng để thesis có tính khoa học và không overclaim.",
                "Hướng phát triển tiếp theo gồm: tăng episode budget và hard episode mining cho DSCNN-L + PCEN + GE2E; tune EdgeSpotFull T4 + PCEN + Triplet/GE2E; thêm distillation từ teacher mạnh cho compact model; thử SCAF với warmup và loss weight nhỏ; và xây dựng benchmark streaming chính thức cho demo dài.",
            ],
            [],
            [],
        ),
        (
            "Threats to Validity",
            1,
            [
                "Thứ nhất, so sánh với EdgeSpot-4 paper không phải reproduction đầy đủ. Khác biệt có thể đến từ data split, training recipe, implementation detail, augmentation, hardware và checkpoint selection. Vì vậy, so sánh chỉ nên dùng như mốc đối chiếu công khai.",
                "Thứ hai, checkpoint selection dùng GSC-dev 3 runs mỗi 5 epoch, trong khi final test dùng 100 runs. Selection noise vẫn có thể ảnh hưởng đến best checkpoint. Một protocol mạnh hơn có thể dùng nhiều dev runs hơn hoặc chọn theo tổ hợp ACC@1%FAR, AUC và F1.",
                "Thứ ba, cap620 có gần 3 triệu train files nhưng episode budget cố định. Kết quả phản ánh hiệu quả trong ngân sách train hiện tại, không phải upper bound của toàn bộ dataset.",
                "Thứ tư, SCAF collapse có thể là do hyperparameter hiện tại chứ không phủ định hoàn toàn ý tưởng angular margin. Kết luận đúng là SCAF chưa ổn định trong setting cap620 hiện tại.",
                "Thứ năm, các kết quả Microset/Top500/manifest20/manifest50 có giá trị bối cảnh nhưng không cùng protocol với cap620 fixed. Khi trình bày ranking final, không được trộn chúng như một bảng duy nhất.",
            ],
            [],
            [],
        ),
        (
            "Phụ lục A. Reproducibility Checklist",
            1,
            [
                "Nguồn số liệu chính: `results/cap620_16_pipeline_metrics_long.csv`, `results/cap620_16_pipeline_test100_summary.md`, Colab run id `colab_mswc_cap620_flac_16pipe_e40_ep150_20260611_154517`.",
                "Script protocol: `colab/run_mswc_cap620_16_pipeline_e40_fixed.sh`. Script này hard-code data profile cap620 FLAC, 40 epoch, 150 episode/epoch, 30 class × 10 sample, checkpoint selection theo GSC-dev ACC@1%FAR, final eval dev30/test100.",
                "Lệnh Colab đã dùng: `MAX_SECONDS=172800 SYNC_SECONDS=300 bash colab/run_mswc_cap620_16_pipeline_e40_fixed.sh` trong thư mục `/content/DoAnTotNghiep`.",
                "Không sync audio clips lên Drive. Chỉ sync checkpoints, results, reports, logs_colab, configs, colab và split manifests. Khi Colab báo gần đầy disk, cần dừng duplicate run và dọn `/content` local, không xóa Drive artifact.",
                "Để tái sinh thesis này, chạy: `python scripts/make_final_thesis_vi_2026_06_12.py` từ root project.",
            ],
            [full_tables[2]],
            [],
        ),
        (
            "Tài liệu tham khảo",
            1,
            [
                "Warden, P. Speech Commands: A Dataset for Limited-Vocabulary Speech Recognition. arXiv:1804.03209.",
                "Wang et al. Trainable Frontend for Robust and Far-Field Keyword Spotting / PCEN-related work. arXiv:1607.05666.",
                "Wan et al. Generalized End-to-End Loss for Speaker Verification. arXiv:1710.10467.",
                "Deng et al. Sub-center ArcFace: Boosting Face Recognition by Large-Scale Noisy Web Faces. ECCV 2020 / arXiv:2007.12680.",
                "EdgeSpot: Efficient and High-Performance Few-Shot Model for Keyword Spotting. arXiv:2601.16316.",
                "Project evidence files: `results/cap620_16_pipeline_metrics_long.csv`, `reports/server_far_metrics/server_far_metrics_summary.md`, `reports/microset/result_table.md`, `src/evaluation/protocols.py`, `scripts/train.py`, `scripts/evaluate.py`.",
            ],
            [],
            [],
        ),
    ]

    title_md = "# Few-Shot Open-Set Keyword Spotting dựa trên Embedding và Prototype Inference\n\n"
    title_md += "Bản thesis tiếng Việt - draft ngày 2026-06-12. Cần cập nhật thông tin bìa theo mẫu chính thức của trường trước khi nộp.\n\n"
    title_md += f"Evidence chính: `{CAP620_CSV.relative_to(ROOT)}` và run `{RUN_ID}`.\n\n"
    return title_md, summary_tables + full_tables, figures, sections


def build_markdown(df: pd.DataFrame) -> str:
    title_md, _, _, sections = build_content(df)
    parts = [title_md]
    parts.append("## Mục lục gợi ý\n")
    parts.append(
        "1. Giới thiệu\n2. Cơ sở lý thuyết và công trình liên quan\n3. Thiết kế hệ thống và phương pháp\n4. Thực nghiệm\n5. Kết quả và thảo luận\n6. So sánh với EdgeSpot-4 paper\n7. Demo system và triển khai\n8. Kết luận và hướng phát triển\n\n"
    )

    for title, level, paragraphs, tables, figures in sections:
        parts.append(section(title, level, paragraphs))
        for table in tables:
            parts.append(md_table(table) + "\n")
        for fig in figures:
            rel = fig.path.relative_to(OUT_DIR).as_posix()
            parts.append(f"![{fig.caption}]({rel})\n\n*{fig.caption}*\n")
    return "\n".join(parts)


def build_docx(df: pd.DataFrame) -> None:
    _, _, _, sections = build_content(df)
    document = Document()
    set_default_styles(document)
    set_page_layout(document)
    add_cover(document)
    add_toc_placeholder(document)
    add_page_numbers(document)

    for title, level, paragraphs, tables, figures in sections:
        add_section_docx(document, title, level, paragraphs)
        for table in tables:
            add_table_docx(document, table)
        for fig in figures:
            add_figure_docx(document, fig)
        if level == 1 and title.startswith("Chương"):
            document.add_paragraph("")

    document.add_section(WD_SECTION_START.NEW_PAGE)
    document.add_heading("Phụ lục B. Bảng metric đầy đủ", level=1)
    full_far1 = make_metric_table(df, "test100_far1", "Bảng phụ lục B1. Toàn bộ 16 pipeline trên GSC-test100@1%FAR.")
    full_far5 = make_metric_table(df, "test100_far5", "Bảng phụ lục B2. Toàn bộ 16 pipeline trên GSC-test100@5%FAR.")
    add_table_docx(document, full_far1)
    add_table_docx(document, full_far5)

    document.save(OUT_DOCX)


def write_audit(df: pd.DataFrame) -> None:
    audit = OUT_DIR / "Do_An_KWS_final_vi_2026_06_12_audit.txt"
    test1 = df[df["eval"] == "test100_far1"]
    best = test1.loc[test1["Open-set ACC@FAR"].idxmax()]
    lines = [
        f"csv_rows={len(df)}",
        f"unique_pipelines={df['run_tag'].nunique()}",
        f"evals={','.join(sorted(df['eval'].unique()))}",
        "all_status_ok="
        + str(
            (df[["train_status", "dev30_far1_status", "test100_far1_status", "test100_far5_status"]] == "ok")
            .all()
            .all()
        ),
        f"best_test100_far1={best['pipeline']} ACC={best['Open-set ACC@FAR']:.2f} AUC={best['AUC']:.2f} EER={best['EER']:.2f} F1={best['F1']:.2f}",
    ]
    audit.write_text("\n".join(lines) + "\n", encoding="utf-8")


def main() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    df = load_cap620()
    OUT_MD.write_text(build_markdown(df), encoding="utf-8")
    build_docx(df)
    write_audit(df)
    print(OUT_MD)
    print(OUT_DOCX)


if __name__ == "__main__":
    main()

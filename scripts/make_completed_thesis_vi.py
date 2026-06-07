from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "docs" / "thesis"
OUT_MD = OUT_DIR / "Do_An_KWS_completed_vi_2026_06_04.md"
OUT_DOCX = OUT_DIR / "Do_An_KWS_completed_vi_2026_06_04.docx"


MICROSET_ROWS = [
    ["DSCNN-L + MFCC + Triplet", "test100", "100", "-", "80.54%", "40.58%", "91.22%", "18.22%", "68.39%", "73.30%"],
    ["EdgeSpotFull T4 + PCEN + SCAF", "test100", "100", "84.64%", "85.21%", "20.61%", "95.69%", "11.89%", "74.52%", "81.92%"],
    ["EdgeSpotFull T4 + PCEN + SCAF+GE2E", "test100", "100", "84.61%", "86.12%", "21.39%", "95.61%", "11.54%", "77.66%", "82.41%"],
]

MATRIX_ROWS = [
    ["DSCNN-L + MFCC + Triplet", "412,896", "5", "69.30%", "67.31%", "57.93%", "44.57%", "43.18%"],
    ["DSCNN-L + MFCC + SCAF", "412,896", "1", "70.72%", "68.26%", "52.03%", "48.85%", "39.00%"],
    ["DSCNN-L + MFCC + GE2E", "412,896", "5", "72.30%", "73.78%", "78.59%", "28.37%", "60.69%"],
    ["DSCNN-L + MFCC + SCAF+GE2E", "412,896", "2", "69.24%", "66.83%", "52.04%", "48.71%", "39.16%"],
    ["DSCNN-L + PCEN + Triplet", "412,900", "5", "72.24%", "72.67%", "73.71%", "33.38%", "54.98%"],
    ["DSCNN-L + PCEN + SCAF", "412,900", "1", "70.02%", "67.67%", "52.32%", "49.15%", "38.74%"],
    ["DSCNN-L + PCEN + GE2E", "412,900", "5", "76.67%", "79.98%", "85.89%", "22.60%", "67.68%"],
    ["DSCNN-L + PCEN + SCAF+GE2E", "412,900", "1", "70.11%", "67.85%", "50.05%", "50.11%", "37.86%"],
    ["EdgeSpotFull T4 + MFCC + Triplet", "130,594", "5", "69.00%", "66.35%", "48.86%", "50.78%", "37.20%"],
    ["EdgeSpotFull T4 + MFCC + SCAF", "130,594", "5", "69.50%", "67.65%", "53.98%", "47.07%", "40.71%"],
    ["EdgeSpotFull T4 + MFCC + GE2E", "130,594", "5", "69.31%", "67.35%", "52.66%", "48.49%", "39.38%"],
    ["EdgeSpotFull T4 + MFCC + SCAF+GE2E", "130,594", "5", "69.15%", "66.94%", "50.52%", "49.99%", "37.96%"],
    ["EdgeSpotFull T4 + PCEN + Triplet", "130,598", "5", "70.76%", "68.81%", "60.18%", "43.41%", "44.36%"],
    ["EdgeSpotFull T4 + PCEN + SCAF", "130,598", "1", "69.52%", "67.35%", "53.59%", "47.57%", "40.26%"],
    ["EdgeSpotFull T4 + PCEN + GE2E", "130,598", "4", "72.94%", "73.35%", "76.68%", "31.30%", "57.29%"],
    ["EdgeSpotFull T4 + PCEN + SCAF+GE2E", "130,598", "3", "69.24%", "67.00%", "50.95%", "49.69%", "38.22%"],
]

SHORTLIST20_ROWS = [
    ["DSCNN-L + PCEN + GE2E", "412,900", "test100", "100", "82.10 +/- 0.87", "86.05 +/- 0.66", "91.57 +/- 0.58", "16.25 +/- 0.86", "75.90 +/- 1.16"],
    ["EdgeSpotFull T4 + PCEN + GE2E", "130,598", "test100", "100", "79.58 +/- 0.91", "83.06 +/- 0.82", "87.22 +/- 0.75", "20.40 +/- 1.01", "70.46 +/- 1.30"],
]

SHORTLIST50_ROWS = [
    ["DSCNN-L + PCEN + GE2E", "412,900", "100", "80.96 +/- 1.16", "84.68 +/- 0.70", "90.45 +/- 0.68", "17.42 +/- 1.08", "74.34 +/- 1.43"],
    ["EdgeSpotFull T4 + PCEN + GE2E", "130,594", "100", "77.14 +/- 0.89", "82.24 +/- 0.74", "87.74 +/- 0.66", "20.19 +/- 0.90", "70.73 +/- 1.16"],
]

FIGURES = [
    ("reports/full_mswc_matrix_analysis/acc1far_ranked_bar.png", "Hình 1. Xếp hạng ACC@1%FAR của 16 pipeline trong Full MSWC phase-1."),
    ("reports/full_mswc_matrix_analysis/det_summary_heatmap.png", "Hình 2. Tổng hợp DET/EER/FRR theo pipeline."),
    ("reports/full_mswc_matrix_analysis/key_effect_delta_bars.png", "Hình 3. Tác động của các thành phần chính trong ablation."),
]


SECTIONS: list[tuple[int, str, list[str]]] = [
    (1, "Acknowledgements", [
        "Em xin bày tỏ lòng biết ơn chân thành tới thầy hướng dẫn, Dr. Tran Hoang Tung, vì sự hướng dẫn, phản hồi chuyên môn và sự hỗ trợ liên tục trong suốt quá trình thực tập và thực hiện đồ án.",
        "Em cũng xin gửi lời cảm ơn tới Dr. Tran Giang Son vì đã hỗ trợ quyền truy cập máy chủ ICTLab. Nguồn tài nguyên tính toán này đóng vai trò quan trọng trong quá trình tải dữ liệu, huấn luyện mô hình và đánh giá các thí nghiệm của đồ án.",
        "Em xin cảm ơn gia đình vì sự động viên và hỗ trợ tinh thần trong suốt quá trình làm việc. Cuối cùng, em xin cảm ơn bạn bè đã hỗ trợ góp ý và chỉnh sửa để bản báo cáo này được hoàn thiện hơn.",
    ]),
    (1, "I. Introduction", []),
    (2, "1.1. Abstract", [
        "Đồ án này nghiên cứu bài toán few-shot open-set keyword spotting, trong đó hệ thống cần nhận diện các từ khóa mới chỉ từ một số lượng nhỏ mẫu enrollment, đồng thời phải từ chối các âm thanh không thuộc tập từ khóa đã đăng ký. Khác với keyword spotting closed-set truyền thống, bài toán này không chỉ yêu cầu phân loại đúng keyword mà còn phải kiểm soát false accept rate để hạn chế trường hợp unknown speech bị nhận nhầm thành keyword.",
        "Hệ thống được xây dựng theo hướng embedding-based keyword spotting. Audio đầu vào được chuẩn hóa, trích xuất đặc trưng MFCC hoặc mel-PCEN, sau đó đưa qua encoder để tạo embedding. Trong giai đoạn inference, mỗi keyword được biểu diễn bằng prototype tính từ trung bình embedding của các mẫu support. Query audio được so khớp với các prototype bằng khoảng cách L2 và được chấp nhận hoặc từ chối dựa trên ngưỡng open-set.",
        "Đồ án đánh giá nhiều thành phần trong pipeline, bao gồm hai backbone DSCNN-L và EdgeSpotFull T4, hai audio frontend MFCC và PCEN, cùng bốn hướng huấn luyện Triplet, SCAF, GE2E và SCAF+GE2E. Kết quả Microset cho thấy EdgeSpotFull T4 kết hợp PCEN và SCAF+GE2E đạt ACC@5%FAR = 86.12%, AUC = 95.61%, EER = 11.54% và F1 = 82.41% trên GSC test100. Trong phần mở rộng Full MSWC, thí nghiệm 16 pipeline cho thấy PCEN và GE2E là hai thành phần có ảnh hưởng tích cực rõ nhất, đặc biệt với DSCNN-L + PCEN + GE2E.",
        "Đóng góp chính của đồ án là xây dựng và đánh giá một pipeline few-shot open-set KWS có khả năng enroll từ khóa mới, so sánh có hệ thống các kết hợp model-feature-loss, và phát triển demo web hỗ trợ enrollment, single detection, long-audio analysis, open-set testing và calibration.",
    ]),
    (2, "1.2. Context and Motivation", [
        "Keyword Spotting (KWS) là bài toán phát hiện một hoặc nhiều từ khóa mục tiêu trong tín hiệu âm thanh. KWS là thành phần quan trọng trong các hệ thống giao tiếp bằng giọng nói như trợ lý ảo, smart home, điều khiển thiết bị rảnh tay và wake-word detection. Trong các ứng dụng này, hệ thống thường phải hoạt động với độ trễ thấp, tài nguyên hạn chế và cần tránh kích hoạt sai khi người dùng không nói từ khóa.",
        "Nhiều hệ thống KWS truyền thống được thiết kế theo bài toán closed-set classification. Mô hình được huấn luyện trên một tập keyword cố định và khi nhận audio đầu vào, nó sẽ chọn một nhãn trong các lớp đã biết. Cách tiếp cận này phù hợp khi tập từ khóa không thay đổi, nhưng chưa đủ linh hoạt cho trường hợp người dùng muốn thêm từ khóa cá nhân hóa chỉ bằng vài mẫu giọng nói.",
        "Few-shot keyword spotting giải quyết vấn đề này bằng cách học một không gian embedding, trong đó các mẫu cùng keyword nằm gần nhau và các keyword khác nhau nằm xa nhau. Khi người dùng enroll một từ khóa mới, hệ thống chỉ cần tính prototype từ một số mẫu support thay vì huấn luyện lại toàn bộ classifier. Tuy nhiên, trong môi trường thực tế, query audio không phải lúc nào cũng thuộc các keyword đã enroll. Vì vậy, hệ thống cần có khả năng open-set rejection, tức là trả về unknown khi audio không đủ gần với bất kỳ prototype nào.",
        "Động lực của đồ án là xây dựng một pipeline KWS có thể vừa hỗ trợ keyword cá nhân hóa với ít mẫu, vừa giảm false accept trong open-set setting. Ngoài ra, đồ án cũng quan tâm đến khả năng triển khai trên thiết bị tài nguyên hạn chế, nên cần so sánh giữa mô hình có độ chính xác cao và mô hình nhỏ gọn phù hợp với edge/device.",
    ]),
    (2, "1.3. Project Objective", [
        "Mục tiêu thứ nhất của đồ án là xây dựng một pipeline few-shot open-set KWS end-to-end, bao gồm xử lý audio, trích xuất đặc trưng, encoder embedding, prototype inference, threshold calibration và demo web.",
        "Mục tiêu thứ hai là đánh giá ảnh hưởng của ba trục thiết kế chính: model architecture, audio frontend và training loss. Cụ thể, đồ án so sánh DSCNN-L với EdgeSpotFull T4, MFCC với PCEN, và các loss Triplet, SCAF, GE2E, SCAF+GE2E.",
        "Mục tiêu thứ ba là đánh giá hệ thống theo protocol GSC few-shot open-set với các metric phù hợp cho open-set KWS, bao gồm ACC@1%FAR, ACC@5%FAR, FRR@FAR, AUC, EER, F1 và DET curve. Các metric này phản ánh tốt hơn trade-off giữa nhận đúng keyword và từ chối unknown so với accuracy closed-set thông thường.",
        "Mục tiêu cuối cùng là xác định hướng mô hình phù hợp cho hai mục tiêu khác nhau: một hướng ưu tiên accuracy và một hướng ưu tiên compact deployment.",
    ]),
    (2, "1.4. Desired Outcomes", [
        "Kết quả mong đợi của đồ án là một hệ thống KWS có thể enroll keyword mới từ một số ít mẫu audio, nhận diện query bằng prototype matching và từ chối unknown bằng threshold/margin policy.",
        "Về mặt nghiên cứu, đồ án kỳ vọng đưa ra phân tích rõ ràng về vai trò của PCEN, GE2E, SCAF và kiến trúc EdgeSpotFull T4 trong few-shot open-set KWS. Kết quả không chỉ nhằm tìm cấu hình có điểm số cao nhất, mà còn nhằm hiểu vì sao một kết hợp tốt hơn hoặc kém hơn trong từng điều kiện dữ liệu.",
        "Về mặt hệ thống, đồ án kỳ vọng tạo ra demo web có thể minh họa các bước enrollment, detection, long-audio analysis, open-set testing và calibration. Demo này giúp quan sát trực tiếp distance, threshold, margin, top candidates và lý do một sample được accept hoặc reject.",
    ]),
    (2, "1.5. Report Structure", [
        "Phần II trình bày các dataset và pipeline xử lý dữ liệu, bao gồm GSC v2, MSWC English, DEMAND noise, preprocessing và augmentation. Phần III mô tả các thành phần mô hình và loss được dùng trong đồ án. Phần IV trình bày system pipeline, thiết kế thí nghiệm, inference và protocol đánh giá open-set. Phần V báo cáo kết quả thực nghiệm và phân tích các kết hợp model-feature-loss. Phần VI mô tả demo system. Phần VII tổng kết đóng góp, hạn chế và hướng phát triển tiếp theo.",
    ]),
    (1, "II. Dataset", []),
    (2, "2.1. GSC v2 Dataset", [
        "Google Speech Commands v2 (GSC v2) là dataset gồm các audio ngắn khoảng một giây, chứa nhiều command words phổ biến như yes, no, up, down, left, right, on, off, stop và go. Trong đồ án này, GSC v2 không được dùng làm tập train chính cho các mô hình MSWC, mà được dùng làm tập đánh giá cross-dataset theo protocol few-shot open-set.",
        "Lý do chọn GSC cho đánh giá là vì nó phù hợp với ngữ cảnh keyword spotting thực tế: audio ngắn, nhiều speaker, nhiều từ dễ nhầm lẫn và có cả background noise/silence. Protocol gsc_edgespot_exact sử dụng k-shot support để tạo prototype cho mỗi keyword, sau đó đánh giá query trên positive words, negative words và silence. Điều này mô phỏng gần hơn tình huống người dùng enroll một số từ khóa rồi hệ thống phải nhận diện hoặc reject audio mới.",
        "Trong các kết quả báo cáo, GSC-dev được dùng để chọn checkpoint, còn GSC-test100 được dùng để báo cáo kết quả cuối. Test100 nghĩa là đánh giá trung bình trên 100 lần chạy với các episode/few-shot sampling khác nhau, giúp giảm ảnh hưởng của may rủi trong việc chọn support/query samples.",
    ]),
    (2, "2.2. MSWC Dataset", [
        "Multilingual Spoken Words Corpus (MSWC) là dataset gồm các từ đơn được trích từ Common Voice, có nhiều ngôn ngữ, nhiều speaker và quy mô lớn. Trong đồ án, MSWC English là nguồn dữ liệu chính để huấn luyện encoder KWS. Việc dùng MSWC giúp mô hình học embedding từ nhiều từ và nhiều speaker hơn so với chỉ huấn luyện trên một tập command nhỏ.",
        "Đồ án sử dụng nhiều profile MSWC khác nhau. Microset English là profile nhỏ, có official CSV split, được dùng làm evidence chính ban đầu để chọn hướng kiến trúc. Top500 là profile gồm 500 từ phổ biến trong MSWC English, dùng để kiểm tra khả năng mở rộng sang vocabulary lớn hơn. Full MSWC English là profile lớn hơn, gồm hàng chục nghìn từ, được dùng cho ablation và shortlist với manifest cap để kiểm soát chi phí huấn luyện.",
        "Một điểm quan trọng trong pipeline dữ liệu là sử dụng file manifest thay vì quét trực tiếp toàn bộ folder theo keyword. Manifest giúp tránh leakage giữa train/dev/test, kiểm soát số lượng clips mỗi word và đảm bảo các lần chạy có thể reproduce.",
    ]),
    (2, "2.3. DEMAND Noise Dataset", [
        "DEMAND là dataset noise môi trường được dùng cho augmentation trong quá trình huấn luyện. Trong KWS thực tế, audio có thể bị ảnh hưởng bởi tiếng quạt, tiếng phòng, tiếng đường phố hoặc các nguồn nhiễu nền khác. Nếu chỉ huấn luyện trên audio sạch, mô hình có thể hoạt động kém khi đưa vào môi trường thật.",
        "Trong pipeline của đồ án, DEMAND được dùng để trộn noise vào audio huấn luyện với xác suất nhất định. Mục tiêu là giúp encoder học embedding ổn định hơn trước thay đổi về background noise. DEMAND không phải tập đánh giá chính; evaluation vẫn dựa trên GSC few-shot open-set protocol.",
    ]),
    (2, "2.4. Data Preprocessing and Augmentation", [
        "Audio đầu vào được chuẩn hóa về mono, sample rate 16 kHz và độ dài xấp xỉ một giây. Nếu audio ngắn hơn độ dài mục tiêu, hệ thống sẽ pad silence; nếu dài hơn, hệ thống sẽ trim hoặc xử lý theo window tùy pipeline. Cách chuẩn hóa này giúp feature extractor và encoder nhận input có kích thước ổn định.",
        "Sau bước chuẩn hóa waveform, hệ thống trích xuất đặc trưng MFCC hoặc mel spectrogram kết hợp PCEN. MFCC được dùng như baseline truyền thống, còn mel-PCEN được dùng cho các hướng EdgeSpotFull T4 và một số biến thể DSCNN. PCEN giúp chuẩn hóa năng lượng theo kênh, giảm độ nhạy với khác biệt âm lượng và nhiễu nền.",
        "Augmentation gồm noise mixing, time shift và SpecAugment. Noise mixing giúp mô phỏng môi trường âm thanh thực tế. Time shift tạo biến thiên vị trí keyword trong cửa sổ một giây. SpecAugment che một phần theo trục tần số hoặc thời gian trên feature, giúp mô hình không phụ thuộc quá mức vào một vùng phổ cố định.",
    ]),
    (1, "III. Model Architecture and Training Objectives", []),
    (2, "3.1. DSCNN-L Baseline", [
        "DSCNN-L là baseline convolutional neural network nhẹ cho keyword spotting. Mô hình dùng depthwise separable convolution để giảm số tham số so với convolution thông thường, trong khi vẫn giữ khả năng học đặc trưng cục bộ trên trục thời gian-tần số. Trong đồ án, DSCNN-L ban đầu được kết hợp với MFCC và Triplet loss để tạo embedding.",
        "Vai trò của DSCNN-L trong đồ án là baseline và accuracy-oriented candidate. Ở giai đoạn Full MSWC shortlist, DSCNN-L + PCEN + GE2E đạt kết quả tốt nhất về accuracy, dù số tham số lớn hơn EdgeSpotFull T4. Điều này cho thấy DSCNN-L vẫn là lựa chọn mạnh nếu mục tiêu chính là điểm số GSC-test100.",
    ]),
    (2, "3.2. EdgeSpotFull T4 Encoder", [
        "EdgeSpotFull T4 là encoder nhỏ gọn lấy cảm hứng từ hướng EdgeSpot-style KWS. Mục tiêu của mô hình là tạo embedding hiệu quả cho few-shot keyword spotting trong khi vẫn giữ số tham số thấp để phù hợp với edge/device. Trong project, EdgeSpotFull T4 dùng khoảng 130.6k tham số, nhỏ hơn đáng kể so với DSCNN-L khoảng 412.9k tham số.",
        "EdgeSpotFull T4 nhận input dạng mel-PCEN và sinh embedding 64 chiều. Embedding này không trực tiếp đưa vào softmax classifier trong inference, mà được dùng để tính prototype cho từng keyword. Vì vậy, chất lượng embedding space là yếu tố quan trọng nhất: các mẫu cùng keyword cần gần nhau, còn keyword khác nhau và unknown cần tách xa nhau.",
        "Trong Microset, EdgeSpotFull T4 + PCEN + SCAF+GE2E đạt kết quả tốt và được dùng làm hướng compact chính. Tuy nhiên, trên Full MSWC shortlist, DSCNN-L + PCEN + GE2E đạt accuracy cao hơn. Do đó, EdgeSpotFull T4 nên được trình bày như compact edge-oriented candidate, không nên claim là tốt nhất tuyệt đối trong mọi setting.",
    ]),
    (2, "3.3. MFCC and PCEN Frontends", [
        "MFCC là đặc trưng truyền thống trong speech processing, nén thông tin phổ âm thanh theo thang mel và thường được dùng trong các hệ thống nhận dạng tiếng nói/KWS. MFCC có ưu điểm là gọn và ổn định, nhưng có thể mất một phần thông tin chi tiết của phổ khi so với mel spectrogram.",
        "PCEN (Per-Channel Energy Normalization) là phương pháp chuẩn hóa năng lượng theo từng kênh mel, giúp giảm ảnh hưởng của thay đổi âm lượng và nhiễu nền. Trong đồ án, PCEN đặc biệt hiệu quả khi kết hợp với GE2E. Trong Full MSWC phase-1, DSCNN-L + GE2E tăng từ 72.30% ACC@1%FAR với MFCC lên 76.67% với PCEN. EdgeSpotFull T4 + GE2E tăng từ 69.31% lên 72.94% khi đổi từ MFCC sang PCEN.",
    ]),
    (2, "3.4. Triplet Loss", [
        "Triplet loss học embedding bằng bộ ba anchor, positive và negative. Loss này buộc khoảng cách giữa anchor và positive nhỏ hơn khoảng cách giữa anchor và negative một margin nhất định. Đây là loss phổ biến trong metric learning và phù hợp khi muốn học không gian embedding phân biệt class.",
        "Tuy nhiên, trong few-shot KWS của đồ án, inference không dựa trên từng cặp anchor-positive-negative mà dựa trên prototype trung bình của support samples. Vì vậy, Triplet loss không hoàn toàn khớp với cơ chế inference cuối cùng. Kết quả Full MSWC phase-1 cho thấy Triplet thường kém GE2E khi giữ nguyên model và frontend.",
    ]),
    (2, "3.5. SCAF", [
        "SCAF trong đồ án được hiểu là hướng Sub-center ArcFace-style loss. Loss này sử dụng angular margin và sub-centers để tăng độ tách biệt giữa các class trong embedding space. Ý tưởng chính là các mẫu cùng class cần gom lại gần các center tương ứng, trong khi các class khác nhau cần có biên tách rõ ràng.",
        "SCAF có giá trị vì audio keyword có thể có nhiều biến thể do speaker, accent và noise. Sub-center giúp một keyword có thể được biểu diễn bởi nhiều vùng nhỏ trong embedding space thay vì chỉ một center cứng. Tuy nhiên, SCAF đơn lẻ không phải lúc nào cũng cho kết quả tốt nhất trong các thí nghiệm Full MSWC phase-1.",
    ]),
    (2, "3.6. GE2E", [
        "GE2E (Generalized End-to-End) là loss dựa trên centroid/prototype. Trong mỗi episode, GE2E tính centroid cho từng class từ support embeddings và tối ưu để query embedding gần centroid đúng, xa centroid sai. Cơ chế này rất gần với inference few-shot KWS, nơi mỗi keyword được biểu diễn bằng prototype trung bình.",
        "Trong các thí nghiệm Full MSWC, GE2E là loss ổn định nhất. Khi so với Triplet, GE2E cải thiện rõ trên nhiều setting. Ví dụ, với DSCNN-L + PCEN, GE2E đạt 76.67% ACC@1%FAR trong phase-1, cao hơn Triplet 72.24%. Với EdgeSpotFull T4 + PCEN, GE2E đạt 72.94%, cao hơn Triplet 70.76%.",
    ]),
    (2, "3.7. SCAF+GE2E", [
        "SCAF+GE2E là loss hybrid kết hợp hai mục tiêu: SCAF tăng separation theo angular margin, còn GE2E làm training sát hơn với prototype inference. Trên Microset, EdgeSpotFull T4 + PCEN + SCAF+GE2E là cấu hình tốt nhất trong các chỉ số quan trọng như ACC@5%FAR, Keyword ACC và F1.",
        "Tuy nhiên, kết quả Full MSWC phase-1 cho thấy SCAF+GE2E không luôn tốt hơn GE2E đơn lẻ. Điều này không có nghĩa hybrid loss là sai về bản chất; nó cho thấy trọng số loss, learning rate, số epoch và kích thước dữ liệu cần được tune thêm. Vì vậy, trong thesis nên trình bày SCAF+GE2E là hướng hiệu quả trên Microset, còn GE2E là hướng ổn định hơn trong Full MSWC shortlist hiện tại.",
    ]),
    (2, "3.8. Training Protocol", [
        "Training sử dụng episodic sampling. Mỗi episode chọn một số class keyword và một số mẫu trên mỗi class. Thiết kế này mô phỏng few-shot setting vì model liên tục học trong các episode nhỏ gồm nhiều keyword khác nhau. Sau mỗi epoch, model được đánh giá trên validation episode và có thể chạy GSC-dev để chọn checkpoint.",
        "Checkpoint selection trong đồ án ưu tiên GSC-dev theo các chỉ số open-set như ACC@1%FAR hoặc ACC@5%FAR. Sau khi chọn checkpoint, kết quả cuối được báo cáo trên GSC-test100. Cách làm này tránh việc tune trực tiếp trên test set.",
    ]),
    (1, "IV. System Pipeline", []),
    (2, "4.1. Overall Pipeline", [
        "Pipeline tổng quát gồm bốn giai đoạn: data preparation, training, enrollment/inference và evaluation/demo. Ở giai đoạn data preparation, hệ thống tải hoặc đọc dữ liệu MSWC/GSC, chuẩn hóa audio và tạo manifest. Ở giai đoạn training, model học embedding bằng episodic sampler và metric-learning loss. Ở giai đoạn inference, support samples được encode thành prototype và query được so khớp bằng L2 distance. Ở giai đoạn evaluation/demo, hệ thống tính metric open-set và hiển thị kết quả cho người dùng.",
        "Cách thiết kế này tách rõ training và inference. Model không cần biết trước toàn bộ keyword cuối cùng trong demo; nó chỉ cần học embedding tốt. Khi người dùng thêm keyword mới, hệ thống chỉ cần encode vài mẫu support và tính prototype.",
    ]),
    (2, "4.2. Experiment Design", [
        "Đồ án đánh giá các pipeline bằng ba trục chính. Trục architecture gồm DSCNN-L và EdgeSpotFull T4. Trục frontend gồm MFCC và PCEN. Trục loss gồm Triplet, SCAF, GE2E và SCAF+GE2E. Sự kết hợp này tạo thành ma trận 16 pipeline trong Full MSWC phase-1.",
        "Mục đích của ma trận không phải chỉ để tìm một dòng có điểm cao nhất, mà để hiểu ảnh hưởng của từng thành phần. Nếu giữ model và loss cố định nhưng đổi MFCC sang PCEN, ta đo được tác động của frontend. Nếu giữ model và frontend cố định nhưng đổi Triplet sang GE2E, ta đo được tác động của loss. Nếu giữ frontend và loss cố định nhưng đổi DSCNN-L sang EdgeSpotFull T4, ta đo được trade-off giữa accuracy và compactness.",
    ]),
    (2, "4.3. Pipeline Combinations", [
        "Các pipeline được đánh giá gồm DSCNN-L + MFCC/PCEN với Triplet, SCAF, GE2E, SCAF+GE2E; và EdgeSpotFull T4 + MFCC/PCEN với cùng bốn loss. Ban đầu đồ án tập trung vào 12 pipeline không gồm EdgeSpot + MFCC, sau đó bổ sung thêm EdgeSpot + MFCC để hoàn thiện ma trận 16 pipeline. Việc bổ sung này giúp kết luận về tác động của PCEN công bằng hơn vì cả hai architecture đều có đủ MFCC và PCEN.",
        "Trong phần kết quả, các pipeline này được báo cáo bằng ACC@1%FAR, ACC@5%FAR, AUC, EER, FRR và F1. ACC@1%FAR là operating point nghiêm ngặt hơn, phù hợp khi hệ thống cần hạn chế false accept. ACC@5%FAR mềm hơn và thường cho thấy khả năng nhận keyword tốt hơn khi chấp nhận false accept cao hơn một chút.",
    ]),
    (2, "4.4. Inference and Open-set Decision", [
        "Trong inference, mỗi keyword được enroll bằng một số mẫu support. Encoder biến từng mẫu support thành embedding và prototype của keyword là trung bình các embedding đó. Với query audio, hệ thống tính embedding query và đo khoảng cách L2 tới tất cả prototype. Keyword có khoảng cách nhỏ nhất là top-1 candidate.",
        "Quyết định open-set dựa trên threshold và margin. Nếu khoảng cách top-1 nhỏ hơn threshold, query có thể được accept. Nếu khoảng cách top-1 quá lớn, query bị reject thành unknown. Ngoài ra, margin giữa top-1 và top-2 cũng có thể được dùng để chặn các trường hợp hai keyword quá gần nhau. Điều này quan trọng với các từ ngắn và dễ nhầm như no/go/on hoặc three/tree.",
    ]),
    (2, "4.5. Evaluation Protocol and Metrics", [
        "Evaluation của đồ án không chỉ đo closed-set accuracy, vì trong open-set KWS hệ thống còn phải từ chối unknown speech. Do đó, kết quả chính được báo cáo theo các operating point có kiểm soát false accept rate. FAR (False Accept Rate) đo tỷ lệ unknown hoặc negative audio bị nhận nhầm thành keyword. FRR (False Reject Rate) đo tỷ lệ keyword thật bị reject. Hai chỉ số này tạo thành trade-off chính của open-set KWS.",
        "ACC@1%FAR và ACC@5%FAR là accuracy tại ngưỡng được chọn sao cho FAR xấp xỉ 1% hoặc 5%. ACC@1%FAR là điều kiện nghiêm ngặt, phù hợp với hệ thống cần hạn chế kích hoạt sai; ACC@5%FAR mềm hơn và cho biết khả năng nhận keyword khi chấp nhận nhiều false accept hơn.",
        "AUC đo khả năng tách score giữa positive và negative trên nhiều ngưỡng. EER là điểm tại đó FAR và FRR bằng nhau; EER càng thấp nghĩa là hệ thống cân bằng tốt hơn giữa accept keyword và reject unknown. DET curve biểu diễn quan hệ FAR-FRR trên nhiều ngưỡng, giúp so sánh trực quan các pipeline thay vì chỉ nhìn một operating point duy nhất.",
        "Trong protocol GSC-test100, mỗi kết quả được tính trên nhiều lần chạy few-shot khác nhau. Việc báo cáo mean và standard deviation giúp giảm phụ thuộc vào một lần sampling support/query cụ thể và phù hợp hơn với cách đánh giá nghiên cứu.",
    ]),
    (1, "V. Experimental Results and Discussion", []),
    (2, "5.1. Microset Main Result", [
        "Microset là evidence chính trong giai đoạn đầu để chọn hướng kiến trúc. Kết quả cho thấy baseline DSCNN-L + MFCC + Triplet đạt ACC@5%FAR = 80.54%, AUC = 91.22%, EER = 18.22% và F1 = 73.30% trên GSC-test100. Khi chuyển sang EdgeSpotFull T4 + PCEN + SCAF, hệ thống đạt ACC@5%FAR = 85.21%, AUC = 95.69%, EER = 11.89% và F1 = 81.92%. Khi thêm GE2E vào SCAF, cấu hình EdgeSpotFull T4 + PCEN + SCAF+GE2E đạt ACC@5%FAR = 86.12%, EER = 11.54% và F1 = 82.41%.",
        "So với baseline, cấu hình EdgeSpotFull T4 + PCEN + SCAF+GE2E tăng 5.58 điểm phần trăm ACC@5%FAR, tăng 9.11 điểm phần trăm F1 và giảm 6.68 điểm phần trăm EER. Điều này ủng hộ hướng chuyển từ MFCC/Triplet sang mel-PCEN và loss phù hợp hơn với embedding/prototype inference.",
    ]),
    (2, "5.2. Full MSWC 16-pipeline Ablation", [
        "Full MSWC phase-1 được dùng như thí nghiệm ablation để đánh giá ảnh hưởng của từng thành phần pipeline. Thí nghiệm này dùng manifest cap 20 clips/word, huấn luyện 5 epochs và 150 episodes/epoch, vì vậy không nên xem đây là kết quả final. Giá trị chính của thí nghiệm là giúp shortlist các cấu hình triển vọng.",
        "Kết quả cho thấy DSCNN-L + PCEN + GE2E là pipeline tốt nhất trong phase-1 với ACC@1%FAR = 76.67%, ACC@5%FAR = 79.98%, AUC = 85.89% và EER = 22.60%. Trong nhóm EdgeSpot, cấu hình tốt nhất là EdgeSpotFull T4 + PCEN + GE2E với ACC@1%FAR = 72.94%, ACC@5%FAR = 73.35%, AUC = 76.68% và EER = 31.30%.",
        "Hai kết luận quan trọng có thể rút ra từ ma trận này. Thứ nhất, PCEN thường cải thiện rõ khi kết hợp với GE2E. Thứ hai, GE2E phù hợp hơn Triplet trong setting prototype-based inference. Ngược lại, SCAF+GE2E không luôn tốt hơn GE2E đơn lẻ trong Full MSWC phase-1, cho thấy hybrid loss cần tune thêm trọng số và schedule.",
    ]),
    (2, "5.3. Full MSWC Shortlist Results", [
        "Sau khi phase-1 xác định các cấu hình triển vọng, đồ án tiếp tục chạy shortlist dài hơn cho DSCNN-L + PCEN + GE2E và EdgeSpotFull T4 + PCEN + GE2E. Với manifest20, DSCNN-L + PCEN + GE2E đạt ACC@1%FAR = 82.10 +/- 0.87 và ACC@5%FAR = 86.05 +/- 0.66 trên GSC-test100. EdgeSpotFull T4 + PCEN + GE2E đạt ACC@1%FAR = 79.58 +/- 0.91 và ACC@5%FAR = 83.06 +/- 0.82.",
        "Kết quả này cho thấy DSCNN-L là accuracy-oriented candidate tốt hơn trong shortlist hiện tại. Tuy nhiên, EdgeSpotFull T4 có số tham số nhỏ hơn đáng kể, khoảng 130.6k so với 412.9k của DSCNN-L. Vì vậy EdgeSpotFull T4 vẫn có giá trị nếu mục tiêu là compact deployment trên edge/device.",
        "Với manifest50, cả hai mô hình đều giảm điểm so với manifest20. DSCNN-L đạt ACC@1%FAR = 80.96 +/- 1.16, còn EdgeSpotFull T4 đạt 77.14 +/- 0.89. Điều này cho thấy manifest50 là setting khó hơn, thêm nhiều biến thiên acoustic/word hơn và có thể cần schedule dài hơn hoặc tuning tốt hơn. Không nên claim manifest50 cải thiện accuracy; nên dùng nó như robustness follow-up xác nhận cùng xu hướng ranking.",
    ]),
    (2, "5.4. Top500 Recheck", [
        "Top500 là bước mở rộng từ Microset sang vocabulary lớn hơn. Kết quả đáng tin cậy hiện tại là checkpoint epoch13 của EdgeSpotFull T4 + PCEN + SCAF+GE2E, đã được re-evaluate từ artifact local. Trên GSC-test100, checkpoint này đạt ACC@1%FAR = 85.62%, ACC@5%FAR = 88.79%, AUC = 95.34%, EER = 11.51% và F1 = 82.45%.",
        "Kết quả Top500 epoch13 cho thấy hướng EdgeSpotFull T4 + PCEN + SCAF+GE2E có tín hiệu tốt khi mở rộng dữ liệu. Tuy nhiên, các kết quả epoch25 trước đó chỉ nên được mô tả là historical/logged run nếu chưa có checkpoint/result JSON tương ứng. Trong thesis, cần phân biệt rõ artifact reproducible và kết quả log lịch sử để tránh overclaim.",
    ]),
    (2, "5.5. Discussion", [
        "Từ các thí nghiệm, có thể thấy không có một pipeline tốt nhất cho mọi mục tiêu. Nếu mục tiêu là accuracy trên GSC-test100 trong Full MSWC shortlist, DSCNN-L + PCEN + GE2E là lựa chọn mạnh nhất hiện tại. Nếu mục tiêu là model nhỏ gọn cho edge/device, EdgeSpotFull T4 + PCEN + GE2E hoặc EdgeSpotFull T4 + PCEN + SCAF+GE2E vẫn có giá trị, đặc biệt vì số tham số thấp hơn nhiều.",
        "Sự khác biệt giữa Microset và Full MSWC cũng rất quan trọng. Trên Microset, SCAF+GE2E hoạt động tốt, nhưng trên Full MSWC phase-1, GE2E đơn lẻ lại ổn định hơn. Điều này cho thấy kết quả của loss hybrid phụ thuộc vào dataset, số class, số clip mỗi word và training schedule. Vì vậy, thesis nên trình bày SCAF+GE2E như một hướng có triển vọng, không phải kết luận tuyệt đối.",
        "DET curve và các operating point theo FAR là cần thiết vì open-set KWS không chỉ là bài toán phân loại đúng keyword. Trong thực tế, false accept có thể nguy hiểm hơn một số false reject, đặc biệt với wake-word hoặc command system. Do đó, ACC@1%FAR và ACC@5%FAR giúp đánh giá hệ thống ở các mức kiểm soát false accept cụ thể.",
    ]),
    (1, "VI. Demo System", []),
    (2, "6.1. Demo Overview", [
        "Demo web được xây dựng để minh họa pipeline few-shot open-set KWS ngoài các bảng metric. Người dùng có thể enroll keyword bằng audio mẫu, chạy single detection, kiểm tra long audio, thử open-set rejection và quan sát các thông tin như top candidate, distance, threshold và margin.",
        "Vai trò của demo không phải thay thế evaluation chính thức. Các kết quả open-set sampled trong UI chỉ có giá trị minh họa/debug. Kết quả nghiên cứu chính vẫn phải dựa trên GSC gsc_edgespot_exact dev/test với số runs rõ ràng.",
    ]),
    (2, "6.2. Long-audio Analysis", [
        "Long-audio flow cho phép ghép hoặc upload audio dài hơn một giây, sau đó hệ thống chạy detection theo window hoặc segmentation. UI hiển thị timeline, expected words, detected words, missed cases và lý do miss/reject. Phần này hữu ích để phân tích lỗi thực tế, ví dụ keyword bị lệch timing, bị threshold reject hoặc bị nhầm với từ gần âm.",
    ]),
    (2, "6.3. Open-set Calibration", [
        "Open-set calibration giúp người dùng hiểu trade-off giữa keyword accuracy và unknown rejection. Khi threshold thấp, hệ thống reject nhiều hơn, false accept giảm nhưng false reject tăng. Khi threshold cao, hệ thống accept nhiều hơn, keyword accuracy có thể tăng nhưng unknown rejection giảm. Demo hiển thị các chỉ số này để hỗ trợ chọn policy phù hợp.",
    ]),
    (1, "VII. Conclusion", []),
    (2, "7.1. Conclusion", [
        "Đồ án đã xây dựng và đánh giá một pipeline few-shot open-set keyword spotting dựa trên embedding và prototype inference. Hệ thống có khả năng enroll keyword mới bằng một số mẫu support, sau đó nhận diện hoặc reject query audio dựa trên khoảng cách tới prototype và threshold open-set.",
        "Kết quả Microset cho thấy EdgeSpotFull T4 + PCEN + SCAF+GE2E cải thiện rõ so với baseline DSCNN-L + MFCC + Triplet ở các metric chính như ACC@5%FAR, EER và F1. Kết quả Full MSWC phase-1 và shortlist cho thấy PCEN và GE2E là hai thành phần có ảnh hưởng tích cực nhất. Trong shortlist hiện tại, DSCNN-L + PCEN + GE2E là accuracy-oriented candidate tốt nhất, còn EdgeSpotFull T4 là compact edge-oriented candidate có số tham số nhỏ hơn đáng kể.",
        "Nhìn chung, đồ án không chỉ đưa ra một mô hình duy nhất mà còn phân tích trade-off giữa accuracy, model size và open-set behavior. Đây là cơ sở quan trọng để chọn pipeline phù hợp tùy mục tiêu triển khai.",
    ]),
    (2, "7.2. Limitations", [
        "Giới hạn đầu tiên là Full MSWC chưa được huấn luyện ở chế độ full all-clips cho mọi pipeline do giới hạn thời gian, GPU và I/O. Một số thí nghiệm dùng manifest cap 20 hoặc 50 clips/word, nên cần được hiểu như ablation/shortlist thay vì final full-data result.",
        "Giới hạn thứ hai là đồ án chưa claim reproduction đầy đủ của EdgeSpot paper. Mô hình EdgeSpotFull T4 trong project lấy cảm hứng từ hướng EdgeSpot-style, nhưng khác về dữ liệu, training setup và một số chi tiết triển khai.",
        "Giới hạn thứ ba là demo open-set và long-audio có giá trị minh họa/debug, chưa thay thế benchmark chính thức. Để đánh giá streaming trong thực tế, cần thêm các metric như latency, false alarm per hour và miss rate trên audio liên tục.",
    ]),
    (2, "7.3. Future Work", [
        "Hướng tiếp theo là train dài hơn cho hai shortlist chính: DSCNN-L + PCEN + GE2E và EdgeSpotFull T4 + PCEN + GE2E, đặc biệt trên manifest lớn hơn hoặc Top500 full clips. Việc này giúp kiểm tra liệu EdgeSpotFull T4 có thể thu hẹp khoảng cách accuracy với DSCNN-L khi được training tốt hơn hay không.",
        "Một hướng khác là tune trọng số của SCAF+GE2E, learning rate và scheduler để kiểm tra vì sao hybrid loss tốt trên Microset nhưng chưa tốt trên Full MSWC phase-1. Ngoài ra, có thể thử knowledge distillation từ teacher speech model nếu muốn tiến gần hơn các hướng paper-grade.",
        "Về demo, cần bổ sung benchmark streaming chính thức, tối ưu calibration theo từng model profile và xuất báo cáo tự động gồm DET curve, bảng metric và per-word error analysis.",
    ]),
]


def md_table(headers: list[str], rows: list[list[str]]) -> str:
    out = ["| " + " | ".join(headers) + " |", "| " + " | ".join(["---"] * len(headers)) + " |"]
    out.extend("| " + " | ".join(row) + " |" for row in rows)
    return "\n".join(out)


def build_markdown() -> str:
    parts: list[str] = []
    parts.append("# Few-Shot Open-Set Keyword Spotting\n")
    parts.append("Bản nháp tiếng Việt hoàn thiện dựa trên file `Đồ Án.docx`, file tham khảo `Đồ án (1).docx`, PDF `M1-Phan_Thanh_Binh-KWS_Master.pdf` và các report kết quả trong repo.\n")
    for level, title, paras in SECTIONS:
        marker = "#" * level
        parts.append(f"{marker} {title}\n")
        for para in paras:
            parts.append(para + "\n")
        if title == "5.1. Microset Main Result":
            parts.append(md_table(["Model", "Split", "Runs", "ACC@1%FAR", "ACC@5%FAR", "FRR@5%FAR", "AUC", "EER", "KW-ACC", "F1"], MICROSET_ROWS) + "\n")
        if title == "5.2. Full MSWC 16-pipeline Ablation":
            parts.append(md_table(["Pipeline", "Params", "Best epoch", "ACC@1%FAR", "ACC@5%FAR", "AUC", "EER", "F1"], MATRIX_ROWS) + "\n")
        if title == "5.3. Full MSWC Shortlist Results":
            parts.append("Manifest20 shortlist:\n\n")
            parts.append(md_table(["Pipeline", "Params", "Split", "Runs", "ACC@1%FAR", "ACC@5%FAR", "AUC", "EER", "F1"], SHORTLIST20_ROWS) + "\n")
            parts.append("Manifest50 robustness follow-up:\n\n")
            parts.append(md_table(["Pipeline", "Params", "Runs", "ACC@1%FAR", "ACC@5%FAR", "AUC", "EER", "F1"], SHORTLIST50_ROWS) + "\n")
        if title == "5.5. Discussion":
            for fig_path, caption in FIGURES:
                if (ROOT / fig_path).exists():
                    parts.append(f"![{caption}](../../{fig_path})\n")
    return "\n".join(parts)


def set_default_font(document: Document) -> None:
    styles = document.styles
    styles["Normal"].font.name = "Times New Roman"
    styles["Normal"].font.size = Pt(12)
    for style_name in ["Heading 1", "Heading 2", "Heading 3"]:
        styles[style_name].font.name = "Times New Roman"


def add_table(document: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for idx, header in enumerate(headers):
        table.rows[0].cells[idx].text = header
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = value


def add_caption(document: Document, text: str) -> None:
    p = document.add_paragraph(text)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p.runs:
        run.italic = True


def build_docx() -> None:
    document = Document()
    set_default_font(document)
    title = document.add_heading("Few-Shot Open-Set Keyword Spotting", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = document.add_paragraph("Bản nháp thesis tiếng Việt")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    note = document.add_paragraph(
        "Bản này được viết tiếp từ outline trong file Đồ Án.docx, tham khảo cấu trúc file Đồ án (1).docx và M1-Phan_Thanh_Binh-KWS_Master.pdf. "
        "Các số liệu được lấy từ report local trong repo, chưa ghi đè file gốc của người dùng."
    )
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for level, title_text, paras in SECTIONS:
        document.add_heading(title_text, level=level)
        for para in paras:
            document.add_paragraph(para)
        if title_text == "5.1. Microset Main Result":
            add_caption(document, "Bảng 1. Kết quả Microset trên GSC-test100.")
            add_table(document, ["Model", "Split", "Runs", "ACC@1%FAR", "ACC@5%FAR", "FRR@5%FAR", "AUC", "EER", "KW-ACC", "F1"], MICROSET_ROWS)
        if title_text == "5.2. Full MSWC 16-pipeline Ablation":
            add_caption(document, "Bảng 2. Kết quả phase-1 của 16 pipeline Full MSWC.")
            add_table(document, ["Pipeline", "Params", "Best epoch", "ACC@1%FAR", "ACC@5%FAR", "AUC", "EER", "F1"], MATRIX_ROWS)
        if title_text == "5.3. Full MSWC Shortlist Results":
            add_caption(document, "Bảng 3. Shortlist manifest20 trên GSC-test100.")
            add_table(document, ["Pipeline", "Params", "Split", "Runs", "ACC@1%FAR", "ACC@5%FAR", "AUC", "EER", "F1"], SHORTLIST20_ROWS)
            add_caption(document, "Bảng 4. Shortlist manifest50 robustness follow-up.")
            add_table(document, ["Pipeline", "Params", "Runs", "ACC@1%FAR", "ACC@5%FAR", "AUC", "EER", "F1"], SHORTLIST50_ROWS)
        if title_text == "5.5. Discussion":
            for fig_path, caption in FIGURES:
                path = ROOT / fig_path
                if path.exists():
                    document.add_picture(str(path), width=Inches(5.8))
                    add_caption(document, caption)

    document.add_heading("References and Evidence Notes", level=1)
    for item in [
        "reports/microset/result_table.md",
        "reports/full_mswc_matrix_analysis/matrix_best_epoch_metrics.md",
        "reports/full_mswc_matrix_analysis/det_curve_summary.md",
        "reports/full_mswc_shortlist_manifest20/shortlist_results_summary.md",
        "reports/full_mswc_shortlist_manifest50/shortlist_results_summary.md",
        "reports/top500_full_recheck/raw/edgespot_epoch13_reval_test100/gsc_edgespot_exact_k10_results.json",
    ]:
        document.add_paragraph(item, style=None)
    document.save(OUT_DOCX)


def main() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    OUT_MD.write_text(build_markdown(), encoding="utf-8")
    build_docx()
    print(OUT_MD)
    print(OUT_DOCX)


if __name__ == "__main__":
    main()

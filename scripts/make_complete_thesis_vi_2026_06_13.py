from __future__ import annotations

import csv
import json
from dataclasses import dataclass, field
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "docs" / "thesis"
OUT_MD = OUT_DIR / "Do_An_KWS_thesis_day_du_vi_2026_06_13.md"
OUT_DOCX = OUT_DIR / "Do_An_KWS_thesis_day_du_vi_2026_06_13.docx"

CAP620_CSV = ROOT / "results" / "cap620_16_pipeline_metrics_long.csv"
CAP620_DEV_SUMMARY = ROOT / "docs" / "reports" / "cap620_development_20260612_summary_vi.md"
AUDIT_REPORT = ROOT / "docs" / "reports" / "project_clearance_audit_2026_06_13_vi.md"

FIG_DIR = ROOT / "docs" / "thesis" / "assets_final_2026_06_12"
FIGURES = [
    (FIG_DIR / "cap620_top8_acc1far.png", "Hình 1. Top pipeline theo ACC@1%FAR trong thí nghiệm cap620 fixed."),
    (FIG_DIR / "cap620_acc1far_heatmap.png", "Hình 2. Heatmap ACC@1%FAR theo backbone, frontend và loss."),
    (FIG_DIR / "edgespot4_comparison_acc1far.png", "Hình 3. Ranh giới claim khi so sánh với mốc EdgeSpot-4 paper."),
]

PAPER_EDGESPOT_ACC1 = 82.0
PAPER_EDGESPOT_PARAMS = "128k"
PAPER_EDGESPOT_MACS = "29.4M"


@dataclass
class TableBlock:
    caption: str
    headers: list[str]
    rows: list[list[str]]


@dataclass
class FigureBlock:
    path: Path
    caption: str


@dataclass
class Block:
    kind: str
    text: str = ""
    level: int = 1
    table: TableBlock | None = None
    figure: FigureBlock | None = None
    items: list[str] = field(default_factory=list)


def pct(x: float, digits: int = 2) -> str:
    return f"{100.0 * x:.{digits}f}%"


def fnum(x: object, digits: int = 2) -> str:
    return f"{float(x):.{digits}f}"


def mean_pm(row: dict, metric: str) -> str:
    std = row.get(f"{metric} std", "")
    if std not in ("", None):
        return f"{fnum(row[metric])} +/- {fnum(std)}"
    return fnum(row[metric])


def vn_mean_pm(row: dict, metric: str) -> str:
    return mean_pm(row, metric).replace("+/-", "±")


def model_label(raw: str) -> str:
    return {
        "dscnn": "DSCNN-L",
        "edgespot_full": "EdgeSpotFull T4",
    }.get(raw, raw)


def frontend_label(raw: str) -> str:
    raw = str(raw)
    return "PCEN" if raw.lower() in {"pcen", "mel_pcen"} else raw.upper()


def loss_label(raw: str) -> str:
    return {
        "triplet": "Triplet",
        "scaf": "SCAF",
        "ge2e": "GE2E",
        "scaf_ge2e": "SCAF+GE2E",
    }.get(raw, raw)


def pipeline_label(row: dict) -> str:
    return f"{model_label(row['model_family'])} + {frontend_label(row['frontend'])} + {loss_label(row['loss'])}"


def load_cap620_rows() -> list[dict]:
    if not CAP620_CSV.exists():
        return []
    rows: list[dict] = []
    with CAP620_CSV.open("r", encoding="utf-8-sig", newline="") as f:
        for row in csv.DictReader(f):
            row = dict(row)
            row["pipeline"] = pipeline_label(row)
            rows.append(row)
    return rows


def select_rows(rows: list[dict], eval_name: str) -> list[dict]:
    return [r for r in rows if r.get("eval") == eval_name]


def get_row(rows: list[dict], eval_name: str, model: str, frontend: str, loss: str) -> dict | None:
    for row in rows:
        if (
            row.get("eval") == eval_name
            and row.get("model_family") == model
            and frontend_label(row.get("frontend", "")) == frontend
            and row.get("loss") == loss
        ):
            return row
    return None


def fixed_table(rows: list[dict], eval_name: str, caption: str) -> TableBlock:
    order_model = {"dscnn": 0, "edgespot_full": 1}
    order_frontend = {"MFCC": 0, "PCEN": 1}
    order_loss = {"triplet": 0, "scaf": 1, "ge2e": 2, "scaf_ge2e": 3}
    sub = sorted(
        select_rows(rows, eval_name),
        key=lambda r: (
            order_model.get(r.get("model_family", ""), 9),
            order_frontend.get(frontend_label(r.get("frontend", "")), 9),
            order_loss.get(r.get("loss", ""), 9),
        ),
    )
    return TableBlock(
        caption,
        ["Pipeline", "ACC@FAR", "AUC", "EER", "FRR@FAR", "Keyword ACC", "F1"],
        [
            [
                r["pipeline"],
                vn_mean_pm(r, "Open-set ACC@FAR"),
                vn_mean_pm(r, "AUC"),
                vn_mean_pm(r, "EER"),
                vn_mean_pm(r, "FRR@FAR"),
                vn_mean_pm(r, "Keyword ACC"),
                vn_mean_pm(r, "F1"),
            ]
            for r in sub
        ],
    )


def top_table(rows: list[dict], eval_name: str, n: int, caption: str) -> TableBlock:
    sub = sorted(select_rows(rows, eval_name), key=lambda r: float(r["Open-set ACC@FAR"]), reverse=True)[:n]
    return TableBlock(
        caption,
        ["Rank", "Pipeline", "ACC@FAR", "AUC", "EER", "F1"],
        [
            [str(i), r["pipeline"], vn_mean_pm(r, "Open-set ACC@FAR"), vn_mean_pm(r, "AUC"), vn_mean_pm(r, "EER"), vn_mean_pm(r, "F1")]
            for i, r in enumerate(sub, start=1)
        ],
    )


def delta_table(rows: list[dict]) -> TableBlock:
    pairs = [
        (
            "DSCNN-L + GE2E: PCEN so với MFCC",
            get_row(rows, "test100_far1", "dscnn", "MFCC", "ge2e"),
            get_row(rows, "test100_far1", "dscnn", "PCEN", "ge2e"),
        ),
        (
            "EdgeSpotFull T4 + GE2E: PCEN so với MFCC",
            get_row(rows, "test100_far1", "edgespot_full", "MFCC", "ge2e"),
            get_row(rows, "test100_far1", "edgespot_full", "PCEN", "ge2e"),
        ),
        (
            "DSCNN-L + PCEN: GE2E so với Triplet",
            get_row(rows, "test100_far1", "dscnn", "PCEN", "triplet"),
            get_row(rows, "test100_far1", "dscnn", "PCEN", "ge2e"),
        ),
        (
            "EdgeSpotFull T4 + PCEN: GE2E so với Triplet",
            get_row(rows, "test100_far1", "edgespot_full", "PCEN", "triplet"),
            get_row(rows, "test100_far1", "edgespot_full", "PCEN", "ge2e"),
        ),
        (
            "PCEN+GE2E: DSCNN-L so với EdgeSpotFull T4",
            get_row(rows, "test100_far1", "edgespot_full", "PCEN", "ge2e"),
            get_row(rows, "test100_far1", "dscnn", "PCEN", "ge2e"),
        ),
    ]
    out = []
    for name, base, cand in pairs:
        if not base or not cand:
            continue
        out.append(
            [
                name,
                f"{float(cand['Open-set ACC@FAR']) - float(base['Open-set ACC@FAR']):+.2f}",
                f"{float(cand['AUC']) - float(base['AUC']):+.2f}",
                f"{float(cand['EER']) - float(base['EER']):+.2f}",
                f"{float(cand['F1']) - float(base['F1']):+.2f}",
            ]
        )
    return TableBlock(
        "Bảng 4. Chênh lệch chính trong thí nghiệm fixed cap620 tại FAR=1%.",
        ["So sánh", "Delta ACC@1%FAR", "Delta AUC", "Delta EER", "Delta F1"],
        out,
    )


def load_json_metrics(path: Path) -> dict | None:
    if not path.exists():
        return None
    return json.loads(path.read_text(encoding="utf-8"))


def development_table() -> TableBlock:
    candidates = [
        (
            "DSCNN-L + PCEN + GE2E, ep300 composite",
            ROOT / "results" / "dscnn_pcen_ge2e_accdev_ep300_composite_colab_mswc_cap620_development_20260612_050614" / "test100_far1" / "gsc_edgespot_exact_k10_results.json",
            ROOT / "results" / "dscnn_pcen_ge2e_accdev_ep300_composite_colab_mswc_cap620_development_20260612_050614" / "test100_far5" / "gsc_edgespot_exact_k10_results.json",
            "Best accuracy",
        ),
        (
            "EdgeSpotFull T4 + PCEN + Triplet hard, ep300 composite",
            ROOT / "results" / "edgespot_t4_pcen_triplet_hard_ep300_composite_colab_mswc_cap620_development_20260612_050614" / "test100_far1" / "gsc_edgespot_exact_k10_results.json",
            ROOT / "results" / "edgespot_t4_pcen_triplet_hard_ep300_composite_colab_mswc_cap620_development_20260612_050614" / "test100_far5" / "gsc_edgespot_exact_k10_results.json",
            "Collapse, không dùng làm kết quả chính",
        ),
        (
            "EdgeSpotFull T4 + PCEN + GE2E, ep300 composite",
            ROOT / "results" / "edgespot_t4_pcen_ge2e_ep300_composite_colab_mswc_cap620_development_20260612_050614" / "test100_far1" / "gsc_edgespot_exact_k10_results.json",
            ROOT / "results" / "edgespot_t4_pcen_ge2e_ep300_composite_colab_mswc_cap620_development_20260612_050614" / "test100_far5" / "gsc_edgespot_exact_k10_results.json",
            "Best compact",
        ),
    ]
    rows = []
    for label, far1_path, far5_path, note in candidates:
        m1 = load_json_metrics(far1_path)
        m5 = load_json_metrics(far5_path)
        if m1:
            rows.append(
                [
                    label,
                    "1%",
                    f"{pct(m1['open_set_acc_at_1far'])} ± {pct(m1['open_set_acc_at_1far_std'])}",
                    f"{pct(m1['auc'])} ± {pct(m1['auc_std'])}",
                    f"{pct(m1['eer'])} ± {pct(m1['eer_std'])}",
                    f"{pct(m1['f1'])} ± {pct(m1['f1_std'])}",
                    f"{pct(m1['keyword_acc'])} ± {pct(m1['keyword_acc_std'])}",
                    note,
                ]
            )
        if m5:
            rows.append(
                [
                    label,
                    "5%",
                    f"{pct(m5['open_set_acc_at_5far'])} ± {pct(m5['open_set_acc_at_5far_std'])}",
                    f"{pct(m5['auc'])} ± {pct(m5['auc_std'])}",
                    f"{pct(m5['eer'])} ± {pct(m5['eer_std'])}",
                    f"{pct(m5['f1'])} ± {pct(m5['f1_std'])}",
                    f"{pct(m5['keyword_acc'])} ± {pct(m5['keyword_acc_std'])}",
                    note,
                ]
            )
    if not rows:
        rows = [
            ["DSCNN-L + PCEN + GE2E, ep300 composite", "1%", "86.36 ± 1.29", "95.21 ± 0.45", "11.32 ± 0.78", "82.73 ± 1.11", "92.92 ± 0.87", "Best accuracy"],
            ["EdgeSpotFull T4 + PCEN + GE2E, ep300 composite", "1%", "82.87 ± 1.22", "92.41 ± 0.44", "14.82 ± 0.70", "77.85 ± 0.97", "87.29 ± 1.19", "Best compact"],
        ]
    return TableBlock(
        "Bảng 5. Kết quả development run cap620 ngày 2026-06-12.",
        ["Cấu hình", "FAR", "ACC", "AUC", "EER", "F1", "Keyword ACC", "Ghi chú"],
        rows,
    )


def simple_table(caption: str, headers: list[str], rows: list[list[str]]) -> TableBlock:
    return TableBlock(caption, headers, rows)


def p(text: str) -> Block:
    return Block("p", text=text)


def h(level: int, text: str) -> Block:
    return Block("h", text=text, level=level)


def tbl(table: TableBlock) -> Block:
    return Block("table", table=table)


def fig(path: Path, caption: str) -> Block:
    return Block("figure", figure=FigureBlock(path, caption))


def bullets(items: list[str]) -> Block:
    return Block("bullets", items=items)


def build_blocks() -> list[Block]:
    rows = load_cap620_rows()
    blocks: list[Block] = []

    blocks += [
        h(0, "Few-Shot Open-Set Keyword Spotting dựa trên Embedding và Prototype Inference"),
        p("Bản thesis tiếng Việt đầy đủ, sinh ngày 2026-06-13. Bản này gom lại logic dự án theo hướng nộp đồ án: từ dữ liệu, mô hình, train, evaluation, demo, so sánh paper, hạn chế và hướng phát triển. Các thông tin họ tên, mã sinh viên, khoa, trường và tên giảng viên hướng dẫn cần thay bằng thông tin chính thức trước khi nộp."),
        h(1, "Thông tin báo cáo"),
        tbl(simple_table(
            "Bảng 0. Thông tin cần điền trước khi nộp.",
            ["Trường", "Giá trị"],
            [
                ["Tên đề tài", "Few-Shot Open-Set Keyword Spotting dựa trên Embedding và Prototype Inference"],
                ["Sinh viên", "[Điền họ tên sinh viên]"],
                ["Mã sinh viên", "[Điền mã sinh viên]"],
                ["Giảng viên hướng dẫn", "[Điền tên giảng viên hướng dẫn]"],
                ["Đơn vị", "[Điền khoa/trường]"],
                ["Ngày", "2026-06-13"],
            ],
        )),
        h(1, "Lời cảm ơn"),
        p("Em xin gửi lời cảm ơn tới giảng viên hướng dẫn, các thầy cô và các anh chị đã hỗ trợ về chuyên môn, tài nguyên tính toán và phản hồi kỹ thuật trong quá trình thực hiện đồ án. Dự án này yêu cầu nhiều vòng thử nghiệm trên Google Colab, server và máy local, vì vậy sự hỗ trợ về môi trường chạy, GPU, lưu trữ và định hướng nghiên cứu có vai trò rất quan trọng."),
        p("Em cũng xin cảm ơn gia đình và bạn bè đã động viên trong quá trình làm việc. Bản báo cáo này được viết theo hướng có thể tiếp tục chỉnh sửa theo mẫu hình thức chính thức của nhà trường; các phần bìa, chữ ký, định dạng citation và phụ lục minh chứng cần được hoàn thiện theo yêu cầu trước khi nộp."),
        h(1, "Tóm tắt"),
        p("Đồ án nghiên cứu bài toán few-shot open-set keyword spotting (KWS), trong đó hệ thống phải nhận diện các từ khóa mới chỉ từ một số ít mẫu enrollment và đồng thời từ chối các âm thanh không thuộc tập từ khóa đã đăng ký. Khác với KWS closed-set, bài toán open-set yêu cầu kiểm soát false acceptance rate (FAR), vì một false accept có thể kích hoạt sai lệnh thoại trong hệ thống thực tế."),
        p("Hướng tiếp cận của đồ án là embedding-based KWS. Audio đầu vào được chuẩn hóa về mono 16 kHz, cắt hoặc đệm về khoảng một giây, sau đó trích xuất MFCC hoặc mel-PCEN. Encoder biến feature thành embedding đã chuẩn hóa L2. Khi người dùng enroll một keyword mới, hệ thống tính prototype bằng trung bình embedding của support samples. Query audio được so sánh với các prototype bằng khoảng cách L2; nếu khoảng cách vượt ngưỡng hoặc policy reject, hệ thống trả về unknown."),
        p("Đồ án đánh giá có hệ thống hai backbone DSCNN-L và EdgeSpotFull T4, hai frontend MFCC và PCEN, cùng bốn hướng loss Triplet, SCAF, GE2E và SCAF+GE2E. Thí nghiệm fixed 16-pipeline trên MSWC English cap620 FLAC tạo 48 dòng metric cho dev30@1%FAR, test100@1%FAR và test100@5%FAR. Development run sau đó tăng budget lên 60 epoch x 300 episode và chọn checkpoint theo composite metric gồm ACC@1%FAR, AUC và F1."),
        p("Kết quả mới nhất cho thấy cấu hình accuracy-oriented tốt nhất là DSCNN-L + PCEN + GE2E, đạt 86.36 ± 1.29% ACC@1%FAR, AUC 95.21 ± 0.45%, EER 11.32 ± 0.78% và F1 82.73 ± 1.11% trên GSC test100. Cấu hình compact tốt nhất hiện tại là EdgeSpotFull T4 + PCEN + GE2E, đạt 82.87 ± 1.22% ACC@1%FAR, AUC 92.41 ± 0.44%, EER 14.82 ± 0.70% và F1 77.85 ± 0.97%."),
        p("So với mốc EdgeSpot-4 paper 82.0% ACC@1%FAR, cấu hình DSCNN-L vượt rõ về mean nhưng có kích thước lớn hơn. Cấu hình EdgeSpotFull T4 compact mới nhỉnh hơn mốc 82.0% về mean, nhưng chênh lệch chỉ 0.87 điểm và run này chưa dùng knowledge distillation; do đó claim đúng là mô hình compact của project đã đạt mức cạnh tranh và hơi cao hơn mean công bố dưới protocol của project, không phải reproduction đầy đủ của paper."),
        h(1, "Abstract"),
        p("This thesis studies few-shot open-set keyword spotting, where a system must recognize newly enrolled keywords from a small number of support samples and reject non-enrolled speech. The proposed pipeline uses audio preprocessing, MFCC or mel-PCEN frontends, neural encoders, L2-normalized embeddings, prototype inference, and threshold-based open-set rejection. Experiments compare DSCNN-L and EdgeSpotFull T4, MFCC and PCEN, and Triplet, SCAF, GE2E, and SCAF+GE2E objectives."),
        p("The latest cap620 development run shows that DSCNN-L + PCEN + GE2E reaches 86.36 ± 1.29% ACC@1%FAR on GSC test100, while the compact EdgeSpotFull T4 + PCEN + GE2E reaches 82.87 ± 1.22%. The results indicate that PCEN and GE2E are the most stable contributors for this prototype-based open-set KWS setting. The thesis also analyzes failure cases such as SCAF collapse and hard-triplet collapse, and separates benchmark evidence from demo-level UI calibration."),
        h(1, "Danh mục thuật ngữ viết tắt"),
        tbl(simple_table(
            "Bảng 1. Thuật ngữ viết tắt.",
            ["Ký hiệu", "Ý nghĩa"],
            [
                ["KWS", "Keyword Spotting, bài toán phát hiện từ khóa trong audio"],
                ["GSC", "Google Speech Commands, dataset dùng cho evaluation"],
                ["MSWC", "Multilingual Spoken Words Corpus, dataset dùng cho training"],
                ["MFCC", "Mel-Frequency Cepstral Coefficients"],
                ["PCEN", "Per-Channel Energy Normalization"],
                ["FAR", "False Acceptance Rate, tỷ lệ unknown bị nhận nhầm thành keyword"],
                ["FRR", "False Rejection Rate, tỷ lệ keyword thật bị reject"],
                ["EER", "Equal Error Rate"],
                ["AUC", "Area Under ROC Curve"],
                ["GE2E", "Generalized End-to-End loss"],
                ["SCAF", "Sub-center ArcFace-style loss"],
                ["KD", "Knowledge Distillation"],
            ],
        )),
        h(1, "Mục lục"),
        p("Trong file Word, mục lục tự động được chèn bằng field TOC. Nếu mở bằng Microsoft Word, chọn mục lục rồi nhấn Update Field để Word cập nhật số trang."),
    ]

    blocks += chapter_1()
    blocks += chapter_2()
    blocks += chapter_3()
    blocks += chapter_4()
    blocks += chapter_5(rows)
    blocks += chapter_6()
    blocks += chapter_7()
    blocks += chapter_8()
    blocks += chapter_9()
    blocks += appendices(rows)
    return blocks


def chapter_1() -> list[Block]:
    return [
        h(1, "Chương 1. Giới thiệu"),
        h(2, "1.1. Bối cảnh Keyword Spotting"),
        p("Keyword Spotting là bài toán phát hiện một hoặc nhiều từ khóa trong tín hiệu âm thanh. Trong các hệ thống trợ lý giọng nói, smart home, thiết bị nhúng hoặc điều khiển rảnh tay, KWS thường là tầng đầu tiên quyết định liệu hệ thống có cần phản hồi hay không. Tầng này phải chạy nhanh, ổn định và ít gây kích hoạt sai."),
        p("KWS truyền thống thường được mô hình hóa như bài toán closed-set classification: model chọn một nhãn trong tập keyword cố định. Cách này hiệu quả khi vocabulary không đổi, nhưng không linh hoạt khi người dùng muốn thêm keyword cá nhân hóa chỉ bằng vài mẫu giọng nói. Nếu cứ thêm keyword mới bằng cách train lại classifier, hệ thống trở nên tốn kém và khó triển khai cho người dùng cuối."),
        h(2, "1.2. Bài toán few-shot open-set KWS"),
        p("Few-shot KWS chuyển trọng tâm từ classifier cố định sang embedding space. Encoder được huấn luyện để đưa các mẫu cùng keyword lại gần nhau và đẩy các keyword khác nhau ra xa nhau. Khi thêm một keyword mới, hệ thống chỉ cần encode vài mẫu support và tính prototype. Query audio sau đó được so sánh với prototype gần nhất."),
        p("Open-set là phần khó nhất. Query có thể là keyword đã enroll, một từ gần âm, một từ ngoài vocabulary, silence hoặc noise. Nếu hệ thống luôn chọn prototype gần nhất mà không có ngưỡng reject, mọi unknown đều bị ép thành một keyword nào đó. Vì vậy, trong đồ án này, metric chính không phải closed-set accuracy mà là ACC@FAR, FRR@FAR, AUC, EER và F1."),
        h(2, "1.3. Mục tiêu nghiên cứu"),
        bullets([
            "Xây dựng pipeline few-shot open-set KWS end-to-end từ dữ liệu, model, training, evaluation đến demo UI.",
            "So sánh hai backbone DSCNN-L và EdgeSpotFull T4 để hiểu trade-off giữa accuracy và compactness.",
            "So sánh MFCC và PCEN để đánh giá vai trò của frontend trong cross-dataset evaluation MSWC -> GSC.",
            "So sánh Triplet, SCAF, GE2E và SCAF+GE2E để xác định objective phù hợp với prototype inference.",
            "Đánh giá hệ thống theo gsc_edgespot_exact test100 ở FAR 1% và 5%, tránh overclaim từ demo UI sampled evaluation.",
            "So sánh thận trọng với mốc EdgeSpot-4 paper và chỉ ra phần nào là competitive result, phần nào chưa phải reproduction.",
        ]),
        h(2, "1.4. Câu hỏi nghiên cứu"),
        tbl(simple_table(
            "Bảng 2. Câu hỏi nghiên cứu và loại bằng chứng.",
            ["Câu hỏi", "Bằng chứng cần dùng", "Kết luận hiện tại"],
            [
                ["RQ1: PCEN có tốt hơn MFCC không?", "Fixed 16-pipeline cap620, so sánh cùng backbone/loss", "PCEN tốt hơn rõ khi kết hợp GE2E, đặc biệt với EdgeSpotFull T4."],
                ["RQ2: Loss nào phù hợp nhất với prototype inference?", "Triplet/SCAF/GE2E/SCAF+GE2E trong cùng protocol", "GE2E ổn định nhất trên cap620; SCAF cần tuning riêng."],
                ["RQ3: Backbone nào nên chọn?", "So DSCNN-L và EdgeSpotFull T4 cùng frontend/loss", "DSCNN-L mạnh hơn về accuracy; EdgeSpotFull T4 là hướng compact."],
                ["RQ4: Project đã vượt EdgeSpot-4 chưa?", "GSC test100 ACC@1%FAR và phân tích protocol", "DSCNN-L vượt rõ mean paper nhưng lớn hơn; EdgeSpotFull T4 nhỉnh hơn mean nhưng chưa claim reproduce."],
                ["RQ5: UI threshold/guard có phải evidence chính không?", "Audit backend/UI và open-set sampled calibration", "Không. UI là demo/debug, benchmark test100 mới là evidence chính."],
            ],
        )),
        h(2, "1.5. Đóng góp của đồ án"),
        p("Đóng góp thứ nhất là xây dựng pipeline KWS hoàn chỉnh dựa trên embedding và prototype inference, cho phép enroll keyword mới bằng ít mẫu audio mà không cần train lại classifier cuối. Đóng góp thứ hai là thực hiện thí nghiệm fixed 16-pipeline trên MSWC cap620 FLAC trong cùng điều kiện để so sánh backbone, frontend và loss một cách công bằng. Đóng góp thứ ba là phát triển development run dài hơn với composite checkpoint selection, đưa DSCNN-L + PCEN + GE2E lên 86.36% ACC@1%FAR và EdgeSpotFull T4 + PCEN + GE2E lên 82.87% ACC@1%FAR."),
        p("Đóng góp thứ tư là phân tích failure modes: SCAF/SCAF+GE2E collapse khi dùng cấu hình chưa tune trên vocabulary lớn, và hard-triplet collapse khi mining quá gắt trong development run. Đóng góp thứ năm là xây dựng demo web phục vụ enrollment, single detection, long-audio analysis, open-set testing và calibration, đồng thời phân biệt rõ demo-level evaluation với benchmark thesis."),
    ]


def chapter_2() -> list[Block]:
    return [
        h(1, "Chương 2. Dữ liệu và chuẩn bị dữ liệu"),
        h(2, "2.1. Tổng quan nguồn dữ liệu"),
        p("Dự án dùng hai nguồn dữ liệu chính: MSWC English để train encoder và Google Speech Commands v2 để evaluate cross-dataset. Ngoài ra, DEMAND noise được dùng cho augmentation. Cách tách này quan trọng: model không học trực tiếp các command GSC theo kiểu classifier closed-set, mà học embedding từ MSWC rồi được đánh giá khả năng few-shot transfer sang GSC."),
        tbl(simple_table(
            "Bảng 3. Vai trò của các dataset trong đồ án.",
            ["Dataset", "Vai trò", "Cách dùng"],
            [
                ["MSWC English", "Training chính", "Tạo train/val words, manifest cap20/cap50/cap220/cap620, train encoder bằng episodic sampling."],
                ["GSC v2", "Evaluation chính", "gsc_edgespot_exact dev/test, k-shot support, query positive/negative/silence, test100."],
                ["DEMAND", "Noise augmentation", "Trộn noise vào waveform trong training để tăng robustness."],
                ["Audio demo local", "Kiểm thử demo", "Single detection, long-audio, timing labels, UI debug."],
            ],
        )),
        h(2, "2.2. Google Speech Commands v2"),
        p("GSC v2 gồm các audio ngắn khoảng một giây, chứa các command phổ biến như yes, no, up, down, left, right, on, off, stop và go. Trong đồ án, GSC không phải nguồn train chính cho MSWC runs. Nó được dùng làm evaluation benchmark vì phù hợp trực tiếp với KWS và có split validation/testing rõ ràng."),
        p("Protocol gsc_edgespot_exact dùng 10 command words cộng với silence thật làm positive targets, và 25 spoken words còn lại làm negative/open-set words. Với mỗi run, hệ thống lấy k-shot support cho mỗi target để tạo prototype, sau đó đánh giá query samples. Final report dùng test100, tức trung bình qua 100 repeated episodes để giảm nhiễu do chọn support/query."),
        h(2, "2.3. MSWC English"),
        p("MSWC English là nguồn train lớn, nhiều từ và nhiều speaker. Dữ liệu gốc có thể được tải từ archive `en.tar.gz`, extract ra thư mục `data/mswc_en/clips/<word>/...`, sau đó dùng metadata để tạo train/val splits. Trong các run Colab mới, audio OPUS được chuyển sang FLAC để tiết kiệm disk nhưng vẫn giữ chất lượng tốt hơn WAV không nén."),
        p("Điểm quan trọng là dự án không dùng trực tiếp mọi file một cách tùy tiện. Các file split và manifest được tạo trong `data/mswc_en/splits`. Manifest giúp kiểm soát số clip tối đa mỗi từ (`max_per_word`), đảm bảo train/val không lẫn nhau, và giúp Colab/server có thể resume hoặc tái lập thí nghiệm."),
        h(2, "2.4. Các profile dữ liệu đã dùng"),
        tbl(simple_table(
            "Bảng 4. Các profile dữ liệu trong lịch sử dự án.",
            ["Profile", "Mục đích", "Ghi chú"],
            [
                ["Microset", "Chọn hướng kiến trúc ban đầu", "Từ vựng nhỏ, giúp chạy nhanh và kiểm tra EdgeSpot/SCAF/GE2E."],
                ["Top500", "Mở rộng lên 500 từ", "Có artifact EdgeSpotFull T4 + PCEN + SCAF+GE2E epoch13 đạt kết quả cao, nhưng không cùng protocol với cap620 fixed."],
                ["manifest20/50/220", "Scale-up theo lượng clip", "Dùng để xem bão hòa dữ liệu và so sánh GE2E/KD."],
                ["cap620 FLAC", "Evidence chính hiện tại", "Khoảng 2.99M train files, 52k val files, 37,387 train words, 763 val words."],
            ],
        )),
        h(2, "2.5. Quy trình chuẩn bị MSWC cap620 FLAC"),
        p("Quy trình cap620 FLAC trong Colab gồm các bước cụ thể sau. Một là mount Google Drive để lưu artifact. Hai là cài dependency cần thiết như ffmpeg, rsync, numpy, scipy, soundfile, scikit-learn, matplotlib, tensorboard, fastapi và uvicorn mà không cài lại torch. Ba là chuẩn bị GSC v2 nếu chưa có `testing_list.txt` và `validation_list.txt`. Bốn là tải metadata MSWC, tạo train/val split bằng seed 42, min-clips 1 và val-fraction 0.02. Năm là tải/extract MSWC English với cap tối đa 620 clip mỗi word. Sáu là chuyển OPUS sang FLAC và xóa OPUS để giảm disk. Bảy là build manifest `train_files_cap620_flac.json` và `val_files_cap620_flac.json`. Tám là sync checkpoints, results, reports, logs, configs, colab và split manifests về Drive."),
        p("Trong run thực tế, Colab báo gần đầy disk khi `/content` dùng khoảng 221 GB / 235.68 GB. Cách xử lý đúng là không chạy thêm duplicate run trong runtime đó, kiểm tra artifact đã sync về Drive, rồi disconnect/delete runtime và tạo runtime mới. Không nên xóa bừa trong Drive; chỉ dọn local `/content` khi chắc artifact đã sync."),
        h(2, "2.6. Preprocessing audio"),
        p("Tất cả audio được đưa về mono 16 kHz và độ dài mục tiêu 1 giây, tương ứng 16,000 samples. Nếu audio ngắn hơn, hệ thống pad silence; nếu dài hơn, hệ thống trim hoặc crop vùng active speech tùy pipeline. Việc chuẩn hóa này giúp feature extractor tạo tensor có kích thước cố định, tránh lỗi shape khi train/evaluate."),
        p("MFCC dùng 40 coefficients và input shape khoảng `(47, 10)` cho DSCNN. Mel/PCEN dùng map time-frequency khoảng `(40, 101)`. Với EdgeSpotFull T4, đường thiết kế tự nhiên là mel-PCEN; MFCC chỉ được giữ làm ablation để kiểm tra vai trò của frontend."),
        h(2, "2.7. File manifest và vì sao không dùng folder scan trực tiếp"),
        p("Một điểm rất quan trọng của dự án là dùng manifest JSON thay vì để DataLoader tự quét toàn bộ thư mục audio. Nếu chỉ quét folder, rất dễ gặp ba lỗi: train/val leakage, số lượng clip mỗi word bị lệch mạnh, và mỗi lần chạy có thể lấy tập file khác nhau. Manifest giải quyết cả ba vấn đề bằng cách ghi rõ danh sách file train và val đã được chọn."),
        p("Với cap620 FLAC, hai file quan trọng là `data/mswc_en/splits/train_files_cap620_flac.json` và `data/mswc_en/splits/val_files_cap620_flac.json`. Mỗi phần tử trong manifest trỏ tới một audio file và nhãn word tương ứng. Khi train, `scripts/train.py` nhận `--train-files` và `--val-files` để dùng đúng manifest. Vì vậy, nếu muốn tái lập kết quả, không chỉ cần checkpoint mà còn cần đúng split manifests."),
        tbl(simple_table(
            "Bảng 5. Các file dữ liệu và ý nghĩa.",
            ["Đường dẫn", "Ý nghĩa", "Có cần lưu không?"],
            [
                ["data/mswc_en/metadata/en_word_counts.json", "Số clip theo từng word trong MSWC English", "Có, để giải thích split/cap."],
                ["data/mswc_en/splits/train_words.json", "Danh sách word dùng để train", "Có."],
                ["data/mswc_en/splits/val_words.json", "Danh sách word dùng để validation", "Có."],
                ["data/mswc_en/splits/train_files_cap620_flac.json", "Manifest train cap620 FLAC", "Rất cần."],
                ["data/mswc_en/splits/val_files_cap620_flac.json", "Manifest validation cap620 FLAC", "Rất cần."],
                ["data/gsc_v2/testing_list.txt", "Split test chính thức của GSC", "Có nếu đóng gói full reproducibility."],
                ["data/gsc_v2/validation_list.txt", "Split dev/validation chính thức của GSC", "Có."],
            ],
        )),
        h(2, "2.8. Từng bước xử lý dữ liệu trên Colab"),
        p("Trong Colab, thứ tự đúng là: mount Drive, giải nén code, cài dependency, tải GSC nếu thiếu, chuẩn bị MSWC, chuyển định dạng audio, build manifest, chạy train, chạy eval, sync artifact. Nếu đảo thứ tự, ví dụ chạy train khi manifest chưa đúng hoặc audio còn ở định dạng chưa được DataLoader hỗ trợ, kết quả có thể lỗi hoặc không tái lập được."),
        p("Script `colab/run_mswc_cap620_16_pipeline_e40_fixed.sh` tự động hóa các bước này. Nó có cơ chế `sync_artifacts_loop` để định kỳ copy checkpoints, results, reports, logs, configs, colab scripts và splits lên Drive. Cơ chế này rất quan trọng vì Colab có thể mất session; nếu không sync, checkpoint tốt có thể mất dù log trước đó đã báo kết quả tốt."),
        p("Khi Colab báo gần đầy disk, không nên cố chạy tiếp. Đầy disk có thể làm fail checkpoint save, fail conversion OPUS->FLAC hoặc làm log/result JSON không ghi đủ. Cách xử lý an toàn là kiểm tra Drive đã có run directory, xem `run.log` và `stages.tsv`, sau đó delete runtime cũ và chạy runtime mới. Không xóa Drive artifact nếu chưa chắc đã backup."),
        h(2, "2.9. Augmentation"),
        p("Augmentation gồm noise mixing, gain, time shift, speed perturb và SpecAugment. Noise mixing dùng DEMAND với xác suất noise khoảng 0.5 và SNR trong khoảng 0-10 dB hoặc theo cấu hình. Time shift dịch keyword trong cửa sổ để model không phụ thuộc vị trí tuyệt đối. SpecAugment mask một phần trục thời gian hoặc tần số để tăng robust feature learning."),
        p("Augmentation chỉ được áp dụng trong training, không dùng để làm sai lệch query evaluation. Evaluation phải phản ánh đúng protocol và không được trộn thêm biến đổi làm mất tính so sánh."),
        h(2, "2.10. Rủi ro dữ liệu và cách kiểm soát"),
        bullets([
            "Rủi ro partial extraction: nếu extract MSWC lỗi nhưng code vẫn chạy, model sẽ train trên dataset thiếu. Code hiện đã sửa để extraction error phải fail thật.",
            "Rủi ro train/val leakage: phải dùng split files cố định, không tự shuffle lại mỗi lần.",
            "Rủi ro disk full: cần xóa archive local sau extract hoặc dùng FLAC, đồng thời sync artifact lên Drive.",
            "Rủi ro so sánh không công bằng: không so một run Top500 với một run cap620 như cùng protocol.",
            "Rủi ro mất checkpoint: luôn lưu `best.pt`, `latest.pt`, epoch checkpoints và log selection metric.",
        ]),
    ]


def chapter_3() -> list[Block]:
    return [
        h(1, "Chương 3. Mô hình, feature và hàm mất mát"),
        h(2, "3.1. Kiến trúc tổng thể"),
        p("Hệ thống gồm encoder embedding, không phải classifier cuối cố định. Encoder nhận feature audio và xuất vector embedding. Trong train, embedding được tối ưu bằng metric learning hoặc centroid-based objectives. Trong inference, embedding support được trung bình thành prototype; query được so với prototype bằng L2 distance."),
        p("Ký hiệu một audio query là x. Sau preprocessing và feature extraction, encoder fθ sinh embedding z = fθ(x). Embedding được chuẩn hóa L2 để ||z||2 = 1. Với một keyword c có k support samples, prototype pc được tính bằng trung bình embedding support rồi chuẩn hóa lại. Query được gán tạm thời cho class có khoảng cách nhỏ nhất d(z, pc). Nếu khoảng cách này nhỏ hơn threshold và policy margin cho phép, hệ thống accept; ngược lại trả về unknown."),
        tbl(simple_table(
            "Bảng 6. Các bước tính score trong inference.",
            ["Bước", "Công thức/mô tả", "Ý nghĩa"],
            [
                ["1", "z = normalize(fθ(x))", "Encode query thành embedding chuẩn hóa."],
                ["2", "pc = normalize(mean(zc,1 ... zc,k))", "Tính prototype cho keyword c từ support samples."],
                ["3", "dc = ||z - pc||2", "Khoảng cách L2 từ query tới prototype c."],
                ["4", "c* = argminc dc", "Keyword candidate gần nhất."],
                ["5", "margin = d_second - d_best", "Độ chắc chắn tương đối giữa top-1 và top-2."],
                ["6", "accept nếu d_best <= threshold và margin >= accept_margin", "Quyết định open-set accept/reject."],
            ],
        )),
        h(2, "3.2. DSCNN-L"),
        p("DSCNN-L là baseline accuracy-oriented của đồ án. Model dùng depthwise separable convolution để giảm chi phí so với convolution thường nhưng vẫn học được pattern thời gian-tần số. Trong code `src/models/dscnn.py`, DSCNN-L có thể dùng MFCC input hoặc mel/PCEN input; khi dùng mel-PCEN, model có PCEN trainable trong frontend."),
        p("DSCNN-L có khoảng 412.9k tham số trong các bảng báo cáo. Nó lớn hơn EdgeSpotFull T4 khoảng ba lần, nhưng có capacity tốt hơn trong development run. Vì vậy, trong thesis nên trình bày DSCNN-L + PCEN + GE2E là cấu hình accuracy chính, không phải compact deployment chính."),
        h(2, "3.3. EdgeSpotFull T4"),
        p("EdgeSpotFull T4 là compact encoder lấy cảm hứng từ EdgeSpot-style KWS. Model có khoảng 130.6k tham số, dùng mel-PCEN, temporal blocks và attention/head để tạo embedding 64 chiều. Mục tiêu của EdgeSpotFull T4 là đạt chất lượng cạnh tranh với kích thước nhỏ hơn, phù hợp hơn với edge/device."),
        p("EdgeSpotFull T4 không nên bị đánh giá chỉ theo một metric duy nhất. Trong fixed 16-pipeline, EdgeSpotFull T4 + PCEN + GE2E có ACC@1%FAR cao nhất trong nhóm compact, còn Triplet có AUC/EER/F1 cạnh tranh hơn. Trong development run, EdgeSpotFull T4 + PCEN + GE2E tăng mạnh lên 82.87% ACC@1%FAR và trở thành compact result chính."),
        h(2, "3.4. Vì sao dùng embedding thay vì softmax classifier"),
        p("Nếu dùng softmax classifier truyền thống, số class ở inference phải trùng với số class đã train. Điều này không phù hợp với yêu cầu người dùng tự enroll keyword mới. Embedding/prototype inference giải quyết vấn đề bằng cách tách encoder khỏi classifier cố định. Encoder chỉ cần học không gian representation tốt; keyword mới được thêm bằng prototype chứ không phải thêm neuron classifier và train lại."),
        p("Cách này cũng phù hợp với few-shot vì mỗi keyword chỉ cần một số support samples. Tuy nhiên, điểm yếu là hệ thống phải có cơ chế reject unknown. Nearest-prototype đơn thuần luôn trả về một keyword, vì vậy open-set threshold là thành phần bắt buộc."),
        h(2, "3.5. MFCC"),
        p("MFCC là frontend truyền thống trong speech processing. Nó nén phổ mel thành cepstral coefficients, giúp input gọn và dễ train. Tuy nhiên, với few-shot open-set KWS, MFCC có thể làm mất chi tiết time-frequency và nhạy với khác biệt amplitude/noise giữa MSWC và GSC. Vì vậy, MFCC được dùng làm baseline/ablation."),
        h(2, "3.6. PCEN"),
        p("PCEN chuẩn hóa năng lượng theo từng kênh mel, giúp giảm ảnh hưởng của biến thiên âm lượng và noise nền. Trong bối cảnh support và query có thể đến từ speaker/microphone khác nhau, PCEN giúp khoảng cách embedding ổn định hơn. Kết quả cap620 xác nhận PCEN là thành phần rất quan trọng, đặc biệt khi đi với GE2E."),
        h(2, "3.7. Triplet loss"),
        p("Triplet loss dùng bộ ba anchor, positive và negative. Loss ép khoảng cách anchor-positive nhỏ hơn anchor-negative ít nhất một margin. Triplet phù hợp metric learning, nhưng hiệu quả phụ thuộc rất mạnh vào mining strategy. Negative quá dễ làm gradient yếu; negative quá khó có thể làm training không ổn định."),
        p("Trong development run, EdgeSpotFull T4 + PCEN + Triplet hard collapse: ACC@1%FAR còn 69.10%, AUC 53.40%, EER 47.84% và F1 39.99%. Điều này không phủ định Triplet, nhưng cho thấy hard mining/hard-pair episode seeding quá gắt cần ablation riêng. Với Triplet, semi-hard mining hoặc giảm hard-pair probability thường an toàn hơn."),
        h(2, "3.8. GE2E"),
        p("GE2E tối ưu trực tiếp theo centroid/prototype trong episode. Với mỗi class, model tính centroid rồi đưa query embedding gần centroid đúng và xa centroid sai. Do inference của hệ thống cũng dùng prototype, GE2E khớp tự nhiên với mục tiêu cuối. Đây là lý do GE2E trở thành loss ổn định nhất trong cap620."),
        h(2, "3.9. SCAF và SCAF+GE2E"),
        p("SCAF là hướng Sub-center ArcFace-style loss. Nó dùng angular margin và sub-centers để tăng tách biệt class trong embedding space. Ý tưởng này hợp lý với speech vì cùng một word có nhiều speaker/accents. Tuy nhiên, khi số lớp train lên tới 37k words, classifier head và scale/margin có thể gây gradient không ổn định nếu không tune."),
        p("SCAF+GE2E từng rất tốt ở Microset/Top500, nhưng trong cap620 nhiều hàng bị collapse với AUC khoảng 50%, FRR@FAR 100%, keyword ACC 9.09% hoặc F1 bằng 0. Kết luận đúng là SCAF chưa ổn trong cấu hình cap620 hiện tại, không phải ý tưởng SCAF sai hoàn toàn. Muốn cứu SCAF cần ablation riêng: giảm weight, giảm scale/margin, warmup GE2E trước, và thử subset nhỏ trước khi chạy full cap620."),
        h(2, "3.10. Knowledge distillation"),
        p("Knowledge distillation là hướng để compact EdgeSpotFull T4 tiến gần hoặc vượt EdgeSpot-4 paper. Paper EdgeSpot có sử dụng teacher-guided objective; trong development run mới của project, KD đang tắt (`RUN_KD=0`). Do đó, kết quả 82.87% của EdgeSpotFull T4 + PCEN + GE2E là rất đáng giá nhưng chưa thể claim reproduction đầy đủ của paper. KD nên được trình bày là hướng phát triển cần làm tiếp hoặc experiment phụ nếu có test100 cùng subset."),
        h(2, "3.11. Open-set threshold và margin"),
        p("Threshold trong benchmark và threshold trong demo UI không giống nhau. Trong benchmark `gsc_edgespot_exact`, threshold được chọn từ score distribution để đạt target FAR, ví dụ 1% hoặc 5%. Nghĩa là threshold được xác định bằng negative examples thật. Trong demo UI, threshold có thể là global threshold do người dùng nhập, hoặc per-class threshold ước lượng từ support samples. Nếu không có negative calibration, demo threshold chỉ là heuristic."),
        p("Margin guard cũng là heuristic. Nó reject khi top-1 và top-2 quá sát nhau. Cơ chế này có thể giúp khi hai keyword thực sự dễ nhầm, nhưng không phải bộ nhận diện từ gần âm. Nó phụ thuộc vào tập keyword đã enroll; thêm hoặc bớt một keyword có thể làm margin thay đổi. Vì vậy, trong bản thesis, margin guard được đặt trong chương demo/triển khai, không đặt làm đóng góp benchmark chính."),
        tbl(simple_table(
            "Bảng 7. Phân biệt threshold benchmark và threshold demo.",
            ["Loại threshold", "Nguồn tính", "Dùng cho", "Có dùng claim final không?"],
            [
                ["Benchmark threshold", "Score distribution trên GSC eval, target FAR", "ACC@1%FAR, ACC@5%FAR, FRR@FAR", "Có."],
                ["Global UI threshold", "Người dùng nhập hoặc apply từ calibration sampled", "Demo single/open-set/long audio", "Không, trừ khi ghi rõ demo-level."],
                ["Per-class UI threshold", "Mean/std support distances từng keyword", "Demo experimental", "Không."],
                ["Margin guard", "Top2 distance - top1 distance", "Reject khi top-1/top-2 quá sát", "Không."],
            ],
        )),
    ]


def chapter_4() -> list[Block]:
    return [
        h(1, "Chương 4. Huấn luyện và đánh giá"),
        h(2, "4.1. Episodic training"),
        p("Training dùng episodic sampling. Mỗi episode chọn một số class và một số sample mỗi class. Trong fixed cap620, mỗi episode có 30 class x 10 samples, 150 episodes/epoch và 40 epochs. Như vậy training không phải supervised epoch quét hết 2.99M files; nó là quá trình sample nhiều episodes từ manifest lớn. Điều này cần viết rõ để tránh hiểu nhầm rằng 40 epoch nghĩa là đi qua toàn bộ dataset 40 lần."),
        p("Với các loss khác nhau, cùng một episode được dùng theo cách khác nhau. Triplet sẽ tạo quan hệ anchor-positive-negative trong batch. GE2E sẽ tách embedding theo class, tính centroid và tối ưu query theo centroid. SCAF dùng label class trong episode để tối ưu angular classification head. Hybrid SCAF+GE2E cộng nhiều thành phần loss theo trọng số."),
        tbl(simple_table(
            "Bảng 8. Ý nghĩa các tham số training chính.",
            ["Tham số", "Ý nghĩa", "Giá trị dùng trong fixed cap620"],
            [
                ["epochs", "Số vòng train cấp cao", "40"],
                ["episodes_per_epoch", "Số episode sample trong mỗi epoch", "150"],
                ["n_classes", "Số keyword/class trong một episode", "30"],
                ["n_samples", "Số sample mỗi class trong episode", "10"],
                ["k-shot eval", "Số support samples mỗi keyword khi đánh giá", "10"],
                ["gsc-dev-runs", "Số run GSC-dev dùng khi chọn checkpoint", "3"],
                ["final test runs", "Số run GSC-test cuối", "100"],
            ],
        )),
        h(2, "4.2. Optimizer và scheduler"),
        p("Cấu hình mặc định dùng Adam với learning rate 0.001 và weight decay 0.0001. Scheduler chính là CosineAnnealingWarmRestarts với T_0=10, T_mult=2 và eta_min=1e-5. Gradient clipping 5.0 được dùng để giảm rủi ro gradient explosion, đặc biệt với hybrid loss hoặc hard mining."),
        h(2, "4.3. Checkpoint selection"),
        p("Checkpoint không được chọn theo train loss đơn thuần. Với `--select-by-gsc-dev`, model được evaluate trên GSC-dev theo protocol gsc_edgespot_exact và chọn best checkpoint theo ACC@1%FAR hoặc composite metric. Fixed 16-pipeline dùng GSC-dev ACC@1%FAR mỗi 5 epoch với 3 runs. Development run dùng composite metric bằng trung bình ACC@1%FAR, AUC và F1, giúp tránh trường hợp chỉ tối ưu một operating point nhưng AUC/F1 xấu."),
        p("Trong code `scripts/train.py`, `best.pt` được lưu khi metric selection hiện tại tốt hơn best trước đó. `latest.pt` được cập nhật để resume nếu run bị ngắt. Một số runner còn truyền `--initial-best-metric` khi resume để tránh việc một checkpoint sau resume nhưng tệ hơn ghi đè `best.pt` cũ. Vì vậy, khi dùng artifact, cần ưu tiên `best.pt` nhưng vẫn kiểm tra log selection để biết best theo metric nào."),
        h(2, "4.4. Một epoch train diễn ra như thế nào"),
        p("Ở mức triển khai, một epoch train gồm các bước: DataLoader sinh các episodes từ manifest; waveform được load và chuyển thành feature MFCC/mel; augmentation có thể được áp dụng; encoder tạo embedding; embedding được L2-normalize; loss được tính theo Triplet/SCAF/GE2E/KD; optimizer backprop và update; scheduler cập nhật learning rate; TensorBoard/log ghi loss và các statistic. Sau mỗi `val_every` hoặc `gsc_dev_every`, model được evaluate để quyết định checkpoint."),
        tbl(simple_table(
            "Bảng 9. Vòng đời một checkpoint.",
            ["Artifact", "Khi nào tạo", "Dùng để làm gì"],
            [
                ["epoch_XX.pt", "Theo `--save-every` hoặc mỗi N epoch", "Debug, rollback, xem diễn tiến training."],
                ["latest.pt", "Cuối epoch hoặc theo runner", "Resume khi runtime bị ngắt."],
                ["best.pt", "Khi selection metric tốt nhất", "Dùng cho final eval và demo."],
                ["run.log", "Trong suốt run", "Chứng minh cấu hình, lỗi, metric theo thời gian."],
                ["stages.tsv", "Sau mỗi stage Colab runner", "Xem stage train/eval đã ok hay chưa."],
            ],
        )),
        h(2, "4.5. Fixed 16-pipeline cap620 protocol"),
        tbl(simple_table(
            "Bảng 10. Cấu hình fixed 16-pipeline cap620.",
            ["Trường", "Giá trị"],
            [
                ["Run ID", "colab_mswc_cap620_flac_16pipe_e40_ep150_20260611_154517"],
                ["Dataset", "MSWC English cap620 FLAC"],
                ["Train files", "2,989,780"],
                ["Validation files", "52,399"],
                ["Train words", "37,387"],
                ["Validation words", "763"],
                ["Epochs", "40"],
                ["Episodes/epoch", "150"],
                ["Episode shape", "30 classes x 10 samples"],
                ["Checkpoint selection", "GSC-dev ACC@1%FAR, every 5 epochs, 3 runs"],
                ["Final evaluation", "dev30@1%FAR, test100@1%FAR, test100@5%FAR"],
            ],
        )),
        h(2, "4.6. Ma trận 16 pipeline"),
        tbl(simple_table(
            "Bảng 11. Ma trận 16 pipeline.",
            ["Nhóm", "Backbone", "Frontend", "Loss"],
            [
                ["1", "DSCNN-L", "MFCC", "Triplet / SCAF / GE2E / SCAF+GE2E"],
                ["2", "DSCNN-L", "PCEN", "Triplet / SCAF / GE2E / SCAF+GE2E"],
                ["3", "EdgeSpotFull T4", "MFCC", "Triplet / SCAF / GE2E / SCAF+GE2E"],
                ["4", "EdgeSpotFull T4", "PCEN", "Triplet / SCAF / GE2E / SCAF+GE2E"],
            ],
        )),
        h(2, "4.7. Development run cap620"),
        p("Sau fixed 16-pipeline, development runner chỉ chạy các nhánh có khả năng cải thiện evidence: accuracy branch DSCNN-L + PCEN + GE2E với budget lớn hơn, compact branch EdgeSpotFull T4 + PCEN + Triplet/GE2E, KD optional và SCAF ablation optional. Trong run 2026-06-12, KD và SCAF ablation đều tắt; kết quả chính đến từ accuracy branch và compact GE2E branch."),
        tbl(simple_table(
            "Bảng 12. Cấu hình development run 2026-06-12.",
            ["Trường", "Giá trị"],
            [
                ["Run ID", "colab_mswc_cap620_development_20260612_050614"],
                ["Runtime", "Google Colab A100 40GB"],
                ["RUN_ACCURACY", "1"],
                ["RUN_COMPACT", "1"],
                ["RUN_KD", "0"],
                ["RUN_SCAF_ABLATION", "0"],
                ["Accuracy branch", "DSCNN-L + PCEN + GE2E, 60 epochs x 300 episodes"],
                ["Compact branch", "EdgeSpotFull T4 + PCEN + Triplet hard / GE2E, 60 epochs x 300 episodes"],
                ["Checkpoint selection", "Composite = mean(ACC@1%FAR, AUC, F1) trên GSC-dev"],
                ["Final eval", "dev30_far1, test100_far1, test100_far5"],
            ],
        )),
        h(2, "4.8. Evaluation metrics"),
        p("AUC đo chất lượng ranking của score trên toàn bộ ngưỡng. EER là điểm FAR và FRR bằng nhau; EER càng thấp càng tốt. FRR@FAR cho biết ở một target FAR cố định, bao nhiêu keyword thật bị reject. ACC@FAR là open-set multiclass accuracy tại threshold không vượt target FAR. F1 cân bằng precision và recall, hữu ích khi một mô hình có accuracy nhìn cao nhưng reject quá nhiều positive."),
        p("Trong báo cáo này, ACC@1%FAR là metric chính để so sánh nghiêm ngặt với EdgeSpot-4 paper. ACC@5%FAR là operating point mềm hơn, giúp thấy tiềm năng nhận keyword khi cho phép false accept cao hơn. Khi một pipeline có AUC=50%, EER=50%, FRR=100% và F1=0, pipeline đó bị coi là collapse dù open-set ACC có thể quanh 69.44% do tỷ lệ unknown/negative trong protocol."),
        h(2, "4.9. Cách đọc kết quả test100"),
        p("Mỗi kết quả test100 là trung bình qua 100 runs. Các cột có dạng mean ± std. Ví dụ ACC@1%FAR = 86.36 ± 1.29 nghĩa là trung bình open-set accuracy ở operating point FAR 1% là 86.36%, với độ lệch chuẩn giữa các runs là 1.29 điểm phần trăm. Khi so sánh hai mô hình chênh nhau rất nhỏ, cần nhìn cả std; nếu khoảng sai số chồng lấn, không nên claim thắng tuyệt đối."),
        p("AUC/EER/F1 không phụ thuộc đúng một operating point như ACC@FAR. Vì vậy, nếu một model có ACC@1%FAR cao hơn nhưng AUC/F1 thấp hơn, cần phân tích trade-off threshold. Đây là lý do development run dùng composite metric thay vì chỉ ACC@1%FAR."),
    ]


def chapter_5(rows: list[dict]) -> list[Block]:
    out = [
        h(1, "Chương 5. Kết quả thực nghiệm và phân tích"),
        h(2, "5.1. Kết quả tổng quan fixed cap620"),
        p("Fixed 16-pipeline là evidence ablation sạch nhất vì tất cả pipeline dùng cùng data profile, cùng số epoch, cùng số episode, cùng checkpoint selection và cùng final evaluation. Bảng top pipeline cho thấy DSCNN-L + PCEN + GE2E là cấu hình tốt nhất theo ACC@1%FAR trong fixed run, còn EdgeSpotFull T4 + PCEN + GE2E là compact pipeline tốt nhất theo ACC@1%FAR."),
    ]
    if rows:
        out += [
            tbl(top_table(rows, "test100_far1", 8, "Bảng 9. Top 8 pipeline trong fixed cap620 tại test100 FAR=1%.")),
            tbl(fixed_table(rows, "test100_far1", "Bảng 10. Bảng đầy đủ 16 pipeline tại test100 FAR=1%.")),
            tbl(fixed_table(rows, "test100_far5", "Bảng 11. Bảng đầy đủ 16 pipeline tại test100 FAR=5%.")),
            tbl(delta_table(rows)),
        ]
    else:
        out += [p("Không tìm thấy `results/cap620_16_pipeline_metrics_long.csv`, vì vậy phần bảng fixed cap620 không được sinh tự động trong lần này.")]
    out += [
        h(2, "5.2. Vì sao PCEN tốt hơn MFCC"),
        p("PCEN cải thiện đặc biệt rõ khi kết hợp với GE2E. Lý do là PCEN làm ổn định năng lượng theo kênh mel, giúp embedding bớt nhạy với khác biệt volume, speaker và thiết bị thu. Trong cross-dataset setting MSWC -> GSC, domain shift là vấn đề thật; PCEN giúp distance giữa support và query phản ánh nội dung từ khóa hơn là khác biệt âm lượng."),
        p("Trên fixed cap620, DSCNN-L + GE2E tăng từ 77.08% ACC@1%FAR với MFCC lên 82.34% với PCEN. EdgeSpotFull T4 + GE2E tăng từ 70.76% lên 79.98% khi đổi MFCC sang PCEN. Đây là evidence rất mạnh để chọn PCEN làm frontend mặc định."),
        h(2, "5.3. Vì sao GE2E phù hợp với prototype inference"),
        p("GE2E huấn luyện theo centroid/prototype, trong khi inference cũng dựa trên prototype. Sự khớp giữa objective và inference làm cho GE2E ổn định hơn Triplet trong nhiều setting. Triplet chỉ quan tâm quan hệ từng bộ ba, còn GE2E quan tâm cấu trúc centroid của cả episode. Với few-shot KWS, điều này gần hơn với cách hệ thống được dùng thật."),
        h(2, "5.4. SCAF collapse"),
        p("Nhiều cấu hình SCAF hoặc SCAF+GE2E bị collapse trong cap620. Dấu hiệu là AUC khoảng 50%, EER khoảng 50%, FRR@FAR 100%, keyword ACC 9.09% và F1 bằng 0. Đây không phải kết quả có thể dùng làm model chính. Nguyên nhân hợp lý là số lớp train rất lớn, trong khi SCAF là angular classification-style objective với classifier head lớn; nếu scale, margin hoặc loss weight không phù hợp, gradient có thể làm embedding mất cấu trúc."),
        p("Kết luận khoa học cần viết cẩn thận: SCAF chưa ổn định trong setting cap620 hiện tại, nhưng không bị loại vĩnh viễn. Cần ablation riêng với weight nhỏ hơn, margin/scale thấp hơn, GE2E warmup và subset trước khi chạy full cap620."),
        h(2, "5.5. Development run mới"),
        tbl(development_table()),
        p("Development run cải thiện mạnh so với fixed run. DSCNN-L + PCEN + GE2E tăng từ 82.34 lên 86.36 ACC@1%FAR, AUC tăng từ 92.42 lên 95.21, EER giảm từ 14.89 xuống 11.32 và F1 tăng từ 77.75 lên 82.73. EdgeSpotFull T4 + PCEN + GE2E tăng từ 79.98 lên 82.87 ACC@1%FAR, AUC tăng từ 87.23 lên 92.41, EER giảm từ 20.23 xuống 14.82 và F1 tăng từ 70.68 lên 77.85."),
        p("Hard Triplet trong development run bị collapse, vì vậy không nên dùng kết quả này để kết luận Triplet luôn kém. Nó chỉ cho thấy cấu hình hard mining hiện tại quá gắt. Nếu muốn cứu Triplet, cần chạy lại với semi-hard mining, giảm hard-pair probability hoặc chọn checkpoint theo nhiều metric ổn định hơn."),
    ]
    for path, caption in FIGURES:
        if path.exists():
            out.append(fig(path, caption))
    out += [
        h(2, "5.6. Best accuracy và best compact"),
        p("Best accuracy hiện tại là DSCNN-L + PCEN + GE2E ep300 composite. Cấu hình này nên là result chính khi mục tiêu là đạt điểm cao nhất trên GSC test100. Best compact hiện tại là EdgeSpotFull T4 + PCEN + GE2E ep300 composite. Cấu hình này nên là result chính khi mục tiêu là model nhỏ hơn, gần với hướng EdgeSpot deployment."),
        h(2, "5.7. Không trộn lẫn Microset, Top500 và cap620"),
        p("Microset, Top500 và cap620 có giá trị khác nhau. Microset giúp chọn hướng ban đầu; Top500 cho thấy tiềm năng của EdgeSpot+SCAF+GE2E ở vocabulary vừa; cap620 fixed/development là evidence chính hiện tại cho large-data thesis. Không nên trộn các số từ các profile này vào một bảng ranking duy nhất như thể chúng cùng protocol. Khi đưa vào thesis, mỗi bảng phải ghi rõ data profile, run id, số epoch, số episode, checkpoint selection và evaluation protocol."),
    ]
    return out


def chapter_6() -> list[Block]:
    return [
        h(1, "Chương 6. So sánh với EdgeSpot-4 paper"),
        h(2, "6.1. Mốc so sánh"),
        p("Mốc paper được dùng trong dự án là EdgeSpot-4 đạt 82.0% ACC@1%FAR với khoảng 128k parameters và 29.4M MACs. Đây là mốc quan trọng vì cùng là few-shot KWS ở operating point FAR thấp. Tuy nhiên, project không phải reproduction đầy đủ của paper vì khác code, data recipe, training objective, checkpoint selection và chưa chạy KD trong development run mới."),
        tbl(simple_table(
            "Bảng 12. So sánh với mốc EdgeSpot-4 paper.",
            ["Hệ thống", "Nguồn/profile", "Kích thước", "ACC@1%FAR", "Nhận xét"],
            [
                ["EdgeSpot-4 paper", "Paper EdgeSpot", f"{PAPER_EDGESPOT_PARAMS}, {PAPER_EDGESPOT_MACS}", "82.0%", "Mốc công bố, không phải kết quả chạy lại trong repo."],
                ["DSCNN-L + PCEN + GE2E", "Project cap620 development", "~412.9k params", "86.36 ± 1.29%", "Vượt mean paper rõ, nhưng model lớn hơn."],
                ["EdgeSpotFull T4 + PCEN + GE2E", "Project cap620 development", "~130.6k params", "82.87 ± 1.22%", "Nhỉnh hơn mean paper, nhưng chênh lệch nhỏ và chưa có KD."],
                ["EdgeSpotFull T4 + PCEN + GE2E", "Project cap620 fixed", "~130.6k params", "79.98 ± 0.98%", "Fixed ablation trước development, chưa vượt paper."],
            ],
        )),
        h(2, "6.2. Claim hợp lệ"),
        bullets([
            "Có thể viết: `DSCNN-L + PCEN + GE2E đạt 86.36 ± 1.29% ACC@1%FAR, vượt mốc mean 82.0% của EdgeSpot-4 nhưng dùng model lớn hơn`.",
            "Có thể viết: `EdgeSpotFull T4 + PCEN + GE2E đạt 82.87 ± 1.22% ACC@1%FAR, cạnh tranh và nhỉnh hơn mean EdgeSpot-4 dưới protocol của project`.",
            "Không nên viết: `Project reproduce đầy đủ EdgeSpot-4 paper`, vì chưa cùng recipe và development run chưa bật KD.",
            "Không nên viết: `EdgeSpotFull T4 vượt dứt khoát paper`, vì chênh lệch chỉ 0.87 điểm và nằm trong khoảng sai số.",
        ]),
        h(2, "6.3. Vai trò của KD"),
        p("Nếu mục tiêu là compact model vượt EdgeSpot-4 một cách thuyết phục, hướng hợp lý là thêm KD hoặc teacher-guided objective. Lý do là paper EdgeSpot dùng teacher/distillation để giúp model nhỏ học representation tốt hơn. Trong project, KD đã được chuẩn bị trong runner nhưng development run mới nhất tắt KD để tiết kiệm thời gian và disk. Vì vậy, KD là future work hoặc experiment tiếp theo, không phải claim chính của bản hiện tại."),
    ]


def chapter_7() -> list[Block]:
    return [
        h(1, "Chương 7. Demo system và triển khai"),
        h(2, "7.1. Kiến trúc demo"),
        p("Demo gồm backend FastAPI và frontend React/Vite. Backend tải checkpoint, chọn frontend theo metadata checkpoint, trích xuất embedding, xây enrollment profile, nhận audio upload và trả kết quả detection. Frontend cung cấp giao diện chọn model, enroll GSC/upload, single detection, long-audio analysis, open-set testing, calibration và export report."),
        tbl(simple_table(
            "Bảng 18. Các nhóm endpoint demo chính.",
            ["Nhóm", "Endpoint tiêu biểu", "Chức năng"],
            [
                ["Model", "GET /api/model/profiles, POST /api/model/select", "Liệt kê/chọn checkpoint, clear hoặc rebuild enrollment khi đổi model."],
                ["Enrollment", "POST /api/enroll/gsc, POST /api/enroll/clear", "Enroll keyword từ GSC hoặc xóa enrollment hiện tại."],
                ["Detection", "POST /api/detect/single, POST /api/detect/long", "Nhận diện single audio hoặc long audio."],
                ["Open-set", "POST /api/open-set/test", "Chạy sampled known/unknown GSC test để debug policy."],
                ["Calibration", "POST /api/open-set/calibrate", "Grid-search threshold, margin, per-class và trả best rows."],
                ["Artifacts", "GET /api/artifacts/status", "Kiểm tra file/checkpoint/result phục vụ UI."],
            ],
        )),
        h(2, "7.2. Model switcher"),
        p("Model switcher cho phép chọn checkpoint có sẵn hoặc checkpoint custom. Khi đổi model, enrollment cũ phải rebuild hoặc clear vì embedding space thay đổi. Dùng prototype của model cũ cho model mới là sai về mặt kỹ thuật, vì vector embedding không còn cùng không gian."),
        h(2, "7.3. Enrollment"),
        p("Enrollment nhận nhiều mẫu audio cho mỗi keyword. Backend chuẩn hóa audio, crop active speech, tạo nhiều enrollment views, encode thành exemplars, tính prototype và threshold. Nếu có impostor waveforms, threshold có thể bị giới hạn theo phân bố impostor; nếu không, threshold chủ yếu dựa vào mean/std của support distances."),
        p("Với từng keyword, enrollment profile lưu prototype, exemplars, threshold, mean/std của support distance và quality metadata của audio. Nếu audio quá ngắn, quá im lặng hoặc chất lượng kém, quality check có thể báo lý do. Enrollment profile giúp single detection và long-audio detection dùng cùng logic score."),
        h(2, "7.4. Single detection"),
        p("Single detection encode một audio query, tính khoảng cách đến các keyword profile, lấy top-1/top-2, sau đó quyết định accept/reject theo threshold và margin. Response hiển thị keyword dự đoán, detected/unknown, distance, threshold, margin, confidence và top candidates."),
        h(2, "7.5. Long-audio analysis"),
        p("Long-audio flow chia audio dài thành segments bằng energy hoặc VAD, chạy detection trên từng segment, rồi so với label/timing nếu người dùng cung cấp. UI hiển thị timeline, expected/detected, match/miss/error, lý do reject theo threshold hoặc margin, và bảng chi tiết. Đây là demo engineering hữu ích nhưng chưa phải streaming benchmark chính thức."),
        h(2, "7.6. Open-set sampled evaluation"),
        p("Open-set UI có preset 17 known / 17 unknown trên GSC. Backend chỉ cho candidate_words là các known words đã enroll, còn unknown words phải bị reject. Metric gồm keyword_acc, unknown_reject_acc, false_accept_rate, false_reject_rate, open_set_acc và balanced_score. Đây là công cụ debug/calibration, không thay thế gsc_edgespot_exact test100."),
        p("Calibration grid-search các threshold từ 0.10 đến 1.20, các accept_margin như 0.00, 0.02, 0.05, 0.08, 0.10, và hai lựa chọn per-class true/false. Backend trả ba row: best_balanced, best_open_set và best_keyword. Nếu muốn demo cân bằng, nên apply best_balanced. Nếu muốn hạn chế false accept, xem best_open_set. Nếu muốn nhận keyword nhiều hơn và chấp nhận false accept, xem best_keyword."),
        h(2, "7.7. Audit per-class threshold và close-word guard"),
        p("Per-class threshold đặt ngưỡng riêng cho từng keyword dựa trên độ phân tán support embeddings. Ý tưởng hợp lý, nhưng khi chỉ có ít mẫu enrollment thì mean/std rất nhiễu. Nếu không có negative/impostor calibration, threshold riêng chỉ biết class đó rộng/hẹp trong support set, không biết unknown có dễ chen vào hay không."),
        p("Close-word guard trong code thực chất là margin rejection: `margin = distance(top2) - distance(top1)`. Nó không hiểu phonetic similarity. Nếu unknown gần một prototype và xa prototype thứ hai, margin vẫn lớn và guard không reject. Nếu enroll nhiều từ gần nhau, margin có thể reject cả true positive. Vì vậy guard nên gọi là `reject khi top-1/top-2 quá sát`, không nên gọi là chặn từ gần âm một cách tuyệt đối."),
        tbl(simple_table(
            "Bảng 13. Khuyến nghị UI policy.",
            ["Chức năng", "Trạng thái hiện tại", "Khuyến nghị"],
            [
                ["Per-class threshold", "Có, dựa vào enrollment profile", "Giữ ở Advanced/Experimental, default OFF."],
                ["Close-word guard", "Margin top1-top2, default có thể bật ở một số view", "Đổi tên thành margin rejection, default OFF nếu chưa calibration."],
                ["Open-set calibration", "Có grid search threshold/margin/per-class", "Đưa thành workflow chính: enroll -> calibrate -> apply -> test."],
                ["Demo sampled eval", "Có preset GSC 17/17", "Dùng để debug, không dùng làm claim thesis chính."],
            ],
        )),
        h(2, "7.8. Workflow demo nên trình bày khi bảo vệ"),
        p("Workflow demo nên đi theo thứ tự ngắn gọn và dễ hiểu. Đầu tiên chọn checkpoint best compact hoặc best accuracy. Tiếp theo enroll một tập keyword bằng GSC preset hoặc upload audio. Sau đó chạy open-set calibration và apply row tốt nhất. Cuối cùng chạy single detection hoặc long-audio để minh họa hệ thống accept keyword đúng và reject unknown. Không nên bắt đầu bằng việc chỉnh tay per-class/guard vì sẽ làm người xem tưởng đây là đóng góp chính."),
        tbl(simple_table(
            "Bảng 19. Script demo khuyến nghị.",
            ["Bước", "Thao tác", "Thông điệp cần nói"],
            [
                ["1", "Chọn model profile", "Mỗi checkpoint có embedding space riêng; đổi model phải rebuild/clear enrollment."],
                ["2", "Enroll keyword", "Few-shot: hệ thống chỉ cần vài support samples để tạo prototype."],
                ["3", "Run calibration", "Threshold nên được chọn bằng known/unknown samples, không chọn theo cảm giác."],
                ["4", "Apply calibrated policy", "Policy đang dùng có threshold, per-class, margin rõ ràng."],
                ["5", "Single/long detection", "Hiển thị top candidates, distance, threshold, margin và lý do accept/reject."],
            ],
        )),
        h(2, "7.9. Giới hạn của demo"),
        p("Demo không phải benchmark cuối. Nó phụ thuộc vào số keyword đã enroll, chất lượng mẫu support, seed sample GSC và threshold/margin đang chọn. Nếu demo nhận sai một file, điều đó không nhất thiết phủ định test100; ngược lại, nếu demo chạy tốt trên vài file, cũng không thể dùng để claim model tốt hơn paper. Thesis phải dựa vào benchmark test100, còn demo dùng để chứng minh hệ thống hoạt động end-to-end."),
    ]


def chapter_8() -> list[Block]:
    return [
        h(1, "Chương 8. Thảo luận và threats to validity"),
        h(2, "8.1. Tại sao dự án từng bị rối"),
        p("Dự án bị rối vì có nhiều lớp evidence chạy nối tiếp nhau: Microset, Top500, manifest20/50/220, cap620 fixed, cap620 development, demo UI và server/Colab artifacts. Mỗi lớp có protocol khác nhau. Nếu không tách rõ, người đọc sẽ thấy số nào cũng quan trọng nhưng không biết số nào là final."),
        p("Cách clear đúng là chọn cap620 development làm result chính mới nhất, fixed 16-pipeline làm ablation nền, Microset/Top500 làm evidence lịch sử và phụ trợ, UI làm demo/debug. Mỗi claim phải gắn với run id, dataset, protocol và metric."),
        h(2, "8.2. Threats to validity"),
        bullets([
            "So sánh với EdgeSpot-4 không phải reproduction đầy đủ, vì khác code, split, data recipe, training objective và KD.",
            "Checkpoint selection có noise vì GSC-dev chỉ dùng số runs giới hạn trong lúc train; final test100 ổn định hơn nhưng vẫn phụ thuộc support sampling.",
            "Cap620 có gần 3 triệu train files nhưng training theo episodic budget, không đảm bảo mọi file được quan sát đồng đều.",
            "SCAF collapse có thể do hyperparameter hiện tại, không phủ định hoàn toàn angular margin learning.",
            "Demo UI sampled evaluation phụ thuộc tập từ enrolled, seed, threshold và số mẫu; không thay thế benchmark test100.",
            "Colab disk/quota có thể làm mất session hoặc thiếu artifact nếu không sync thường xuyên.",
        ]),
        h(2, "8.3. Bài học kỹ thuật"),
        p("Bài học lớn nhất là phải tách training objective, evaluation protocol và demo policy. GE2E tốt vì objective gần với inference. PCEN tốt vì giảm domain shift. SCAF cần cẩn thận vì objective classification/angular có thể không ổn với 37k classes nếu dùng weight/scale/margin chưa tune. Demo threshold/guard cần calibration bằng unknown thật, không nên dựa vào cảm giác khi thử vài audio."),
        h(2, "8.4. Claim đúng và claim sai"),
        tbl(simple_table(
            "Bảng 20. Claim nên viết và không nên viết.",
            ["Nên viết", "Không nên viết", "Lý do"],
            [
                ["DSCNN-L + PCEN + GE2E là best accuracy hiện tại.", "Mọi model của project đều tốt hơn paper.", "Chỉ một số cấu hình mạnh; nhiều cấu hình collapse."],
                ["EdgeSpotFull T4 + PCEN + GE2E cạnh tranh và nhỉnh hơn mean EdgeSpot-4.", "Đã reproduce đầy đủ EdgeSpot-4.", "Chưa cùng recipe và chưa bật KD trong development run."],
                ["SCAF cần tuning lại ở cap620.", "SCAF là loss vô dụng.", "SCAF tốt ở Microset/Top500 nhưng collapse ở cấu hình lớn hiện tại."],
                ["UI open-set calibration là demo/debug.", "UI sampled test thay thế test100.", "Benchmark chính là gsc_edgespot_exact test100."],
                ["Hard Triplet collapse trong run hard-mining hiện tại.", "Triplet luôn collapse.", "Triplet phụ thuộc mining; fixed run Triplet vẫn cạnh tranh ở EdgeSpot group."],
            ],
        )),
        h(2, "8.5. Cách kể câu chuyện thesis logic nhất"),
        p("Câu chuyện nên đi theo tuyến: bài toán cần few-shot và open-set; giải pháp là embedding + prototype + threshold; để làm embedding tốt cần chọn frontend, backbone và loss; thí nghiệm 16 pipeline chứng minh PCEN/GE2E ổn định nhất; development run tăng budget và composite selection để cải thiện result; demo minh họa hệ thống end-to-end nhưng không thay thế benchmark; so sánh paper phải thận trọng vì chưa reproduce đầy đủ KD."),
        p("Nếu trình bày theo lịch sử chat hoặc lịch sử chạy Colab, người đọc sẽ thấy rối. Thesis không nên kể mọi lần thử theo thứ tự thời gian; thesis phải gom lại thành các câu hỏi nghiên cứu và dùng các run phù hợp để trả lời từng câu hỏi."),
    ]


def chapter_9() -> list[Block]:
    return [
        h(1, "Chương 9. Kết luận và hướng phát triển"),
        h(2, "9.1. Kết luận"),
        p("Đồ án đã xây dựng pipeline few-shot open-set keyword spotting dựa trên embedding và prototype inference. Hệ thống có thể thêm keyword mới bằng vài support samples, sau đó nhận diện hoặc reject query audio bằng distance threshold. Pipeline được đánh giá bằng protocol open-set phù hợp hơn closed-set accuracy."),
        p("Kết quả mới nhất cho thấy DSCNN-L + PCEN + GE2E là cấu hình accuracy tốt nhất, đạt 86.36 ± 1.29% ACC@1%FAR. EdgeSpotFull T4 + PCEN + GE2E là cấu hình compact tốt nhất, đạt 82.87 ± 1.22% ACC@1%FAR. PCEN và GE2E là hai thành phần ổn định nhất. SCAF/SCAF+GE2E cần tuning riêng trước khi dùng trên vocabulary lớn."),
        h(2, "9.2. Hướng phát triển"),
        bullets([
            "Tiếp tục tối ưu DSCNN-L + PCEN + GE2E bằng tăng episode budget, hard episode mining vừa phải và composite checkpoint selection.",
            "Tối ưu EdgeSpotFull T4 + PCEN + GE2E bằng KD hoặc teacher-guided objective để vượt EdgeSpot-4 thuyết phục hơn.",
            "Chạy SCAF ablation trên subset với scaf_weight nhỏ, scale/margin thấp và GE2E warmup.",
            "Cải tiến Triplet bằng semi-hard mining thay vì hard mining quá gắt.",
            "Chuẩn hóa UI: default global calibrated threshold, per-class/guard ở Advanced, calibration là workflow chính.",
            "Xây dựng streaming benchmark chính thức với latency, false alarm per hour và miss rate trên audio dài.",
        ]),
    ]


def appendices(rows: list[dict]) -> list[Block]:
    out = [
        h(1, "Phụ lục A. Lệnh tái lập"),
        h(2, "A.1. Chạy fixed 16-pipeline trên Colab"),
        p("Lệnh chính đã dùng cho fixed run:"),
        p("MAX_SECONDS=172800 SYNC_SECONDS=300 bash colab/run_mswc_cap620_16_pipeline_e40_fixed.sh"),
        h(2, "A.2. Chạy development run"),
        p("Lệnh development run ưu tiên accuracy và compact, tắt KD/SCAF ablation:"),
        p("MAX_SECONDS=172800 SYNC_SECONDS=300 RUN_ACCURACY=1 RUN_COMPACT=1 RUN_KD=0 RUN_SCAF_ABLATION=0 ACC_EPOCHS=60 ACC_EPISODES=300 COMPACT_EPOCHS=60 COMPACT_EPISODES=300 GSC_SELECT_METRIC=composite bash colab/run_mswc_cap620_development_experiments.sh"),
        h(2, "A.3. Evaluate checkpoint"),
        p("Lệnh evaluation canonical:"),
        p("python scripts/evaluate_edgespot_protocol.py --checkpoint checkpoints/<run_tag>/best.pt --model-family auto --feature-type auto --k-shot 10 --n-runs 100 --gsc-query-split test --output-dir results/<tag>/test100_far1"),
        h(1, "Phụ lục B. Artifact cần lưu"),
        tbl(simple_table(
            "Bảng 14. Artifact quan trọng.",
            ["Loại", "Đường dẫn/ghi chú"],
            [
                ["Best accuracy checkpoint", "checkpoints/dscnn_pcen_ge2e_accdev_ep300_composite_colab_mswc_cap620_development_20260612_050614/best.pt"],
                ["Best compact checkpoint", "checkpoints/edgespot_t4_pcen_ge2e_ep300_composite_colab_mswc_cap620_development_20260612_050614/best.pt"],
                ["Fixed 16-pipeline CSV", "results/cap620_16_pipeline_metrics_long.csv"],
                ["Development result JSON", "results/dscnn_pcen_ge2e_accdev_ep300_composite.../test100_far1/gsc_edgespot_exact_k10_results.json"],
                ["Development compact JSON", "results/edgespot_t4_pcen_ge2e_ep300_composite.../test100_far1/gsc_edgespot_exact_k10_results.json"],
                ["Audit report", "docs/reports/project_clearance_audit_2026_06_13_vi.md"],
            ],
        )),
        h(1, "Phụ lục C. File code chính"),
        tbl(simple_table(
            "Bảng 15. Các file code/thư mục quan trọng.",
            ["File", "Vai trò"],
            [
                ["data/download_mswc.py", "Tải metadata, tạo split, tải/extract MSWC English."],
                ["data/build_mswc_file_splits.py", "Tạo manifest train/val files theo cap và định dạng audio."],
                ["scripts/train.py", "Train encoder với Triplet/SCAF/GE2E/KD và checkpoint selection."],
                ["scripts/evaluate.py", "Evaluation chung nhiều protocol."],
                ["scripts/evaluate_edgespot_protocol.py", "Wrapper chạy gsc_edgespot_exact canonical."],
                ["src/evaluation/protocols.py", "Định nghĩa positive/negative words, episodes và metric aggregation."],
                ["src/models/dscnn.py", "DSCNN-L encoder."],
                ["src/models/edgespot_full.py", "EdgeSpotFull T4 encoder."],
                ["src/models/ge2e.py", "GE2E loss."],
                ["src/models/arcface.py", "ArcFace/Sub-center ArcFace style loss."],
                ["src/demo/api_server.py", "Backend demo, enrollment, detection, open-set test/calibration."],
                ["src/demo/ui/src/App.tsx", "Frontend demo React/Vite."],
            ],
        )),
        h(1, "Phụ lục D. Checklist tái lập kết quả"),
        bullets([
            "Kiểm tra đúng commit/code package dùng cho Colab run.",
            "Kiểm tra `configs/default.yaml` và các env vars trong runner.",
            "Kiểm tra GSC v2 có `testing_list.txt` và `validation_list.txt`.",
            "Kiểm tra manifest cap620 FLAC tồn tại và số file đúng.",
            "Kiểm tra `stages.tsv` báo train/dev30/test100 đều `ok`.",
            "Kiểm tra `best.pt` là checkpoint được chọn theo GSC-dev metric mong muốn.",
            "Kiểm tra result JSON có `n_runs=100`, `target_far=0.01` hoặc `0.05` đúng với bảng.",
            "Không dùng UI sampled result để thay số test100.",
            "Khi copy artifact từ Drive về local, giữ nguyên cấu trúc results/checkpoints/logs.",
        ]),
        h(1, "Phụ lục E. Tài liệu tham khảo gợi ý"),
        tbl(simple_table(
            "Bảng 16. Tài liệu tham khảo cần chuẩn hóa citation trước khi nộp.",
            ["Chủ đề", "Tài liệu gợi ý"],
            [
                ["Google Speech Commands", "Warden, P. Speech Commands: A Dataset for Limited-Vocabulary Speech Recognition. arXiv:1804.03209."],
                ["MSWC", "Multilingual Spoken Words Corpus / MLCommons Spoken Words Corpus documentation and paper."],
                ["PCEN", "Wang et al. Trainable Frontend for Robust and Far-Field Keyword Spotting / PCEN-related work."],
                ["GE2E", "Wan et al. Generalized End-to-End Loss for Speaker Verification. arXiv:1710.10467; GE2E-KWS related work."],
                ["SCAF", "Deng et al. Sub-center ArcFace: Boosting Face Recognition by Large-Scale Noisy Web Faces. ECCV 2020."],
                ["EdgeSpot", "EdgeSpot: Efficient and High-Performance Few-Shot Model for Keyword Spotting. arXiv:2601.16316."],
                ["Triplet loss", "Metric learning / triplet loss foundational papers."],
            ],
        )),
        h(1, "Phụ lục F. Ghi chú về font và file Word"),
        p("File DOCX được sinh bằng python-docx với font mặc định Times New Roman cho Normal, Heading và Table text. Nội dung được ghi trực tiếp bằng Unicode tiếng Việt. Nếu mở trong Word mà mục lục chưa hiện số trang, chọn mục lục và nhấn Update Field. Nếu trường yêu cầu font khác như Arial hoặc Times New Roman cỡ 13, có thể chỉnh style trong Word hoặc sửa script và chạy lại."),
    ]
    return out


def md_escape_table_cell(text: str) -> str:
    return str(text).replace("|", "\\|").replace("\n", " ")


def render_markdown(blocks: Iterable[Block]) -> str:
    lines: list[str] = []
    for block in blocks:
        if block.kind == "h":
            if block.level == 0:
                lines.append(f"# {block.text}")
            else:
                lines.append(f"{'#' * block.level} {block.text}")
            lines.append("")
        elif block.kind == "p":
            lines.append(block.text)
            lines.append("")
        elif block.kind == "bullets":
            for item in block.items:
                lines.append(f"- {item}")
            lines.append("")
        elif block.kind == "table" and block.table:
            t = block.table
            lines.append(f"**{t.caption}**")
            lines.append("")
            lines.append("| " + " | ".join(md_escape_table_cell(h) for h in t.headers) + " |")
            lines.append("| " + " | ".join("---" for _ in t.headers) + " |")
            for row in t.rows:
                lines.append("| " + " | ".join(md_escape_table_cell(c) for c in row) + " |")
            lines.append("")
        elif block.kind == "figure" and block.figure:
            f = block.figure
            if f.path.exists():
                rel = f.path.relative_to(OUT_DIR).as_posix() if f.path.is_relative_to(OUT_DIR) else str(f.path)
                lines.append(f"![{f.caption}]({rel})")
                lines.append("")
                lines.append(f"*{f.caption}*")
                lines.append("")
    return "\n".join(lines).strip() + "\n"


def set_run_font(run, size: float | None = None, bold: bool | None = None, italic: bool | None = None) -> None:
    run.font.name = "Times New Roman"
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for key in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        rfonts.set(qn(key), "Times New Roman")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_paragraph_font(paragraph, size: float = 13) -> None:
    for run in paragraph.runs:
        set_run_font(run, size=size)


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    p0 = cell.paragraphs[0]
    r = p0.add_run(str(text))
    set_run_font(r, size=10.5, bold=bold)
    p0.alignment = WD_ALIGN_PARAGRAPH.LEFT
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def add_toc(paragraph) -> None:
    run = paragraph.add_run()
    fld_char = OxmlElement("w:fldChar")
    fld_char.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = r'TOC \o "1-3" \h \z \u'
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "separate")
    fld_char3 = OxmlElement("w:fldChar")
    fld_char3.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char)
    run._r.append(instr)
    run._r.append(fld_char2)
    run._r.append(fld_char3)


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(2.0)

    styles = doc.styles
    for style_name, size, bold in [
        ("Normal", 13, False),
        ("Heading 1", 16, True),
        ("Heading 2", 14, True),
        ("Heading 3", 13, True),
    ]:
        style = styles[style_name]
        style.font.name = "Times New Roman"
        style.font.size = Pt(size)
        style.font.bold = bold
        rpr = style.element.get_or_add_rPr()
        rfonts = rpr.rFonts
        if rfonts is None:
            rfonts = OxmlElement("w:rFonts")
            rpr.append(rfonts)
        for key in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
            rfonts.set(qn(key), "Times New Roman")


def add_docx_table(doc: Document, table_block: TableBlock) -> None:
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(table_block.caption)
    set_run_font(r, size=11, bold=True)

    table = doc.add_table(rows=1, cols=len(table_block.headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, header in enumerate(table_block.headers):
        set_cell_text(hdr[i], header, bold=True)
        shade_cell(hdr[i], "D9EAF7")
    for row in table_block.rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value)
    doc.add_paragraph()


def build_docx(blocks: Iterable[Block]) -> Document:
    doc = Document()
    configure_document(doc)

    first = True
    for block in blocks:
        if block.kind == "h":
            if block.level == 0:
                p_title = doc.add_paragraph()
                p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p_title.add_run(block.text)
                set_run_font(run, size=18, bold=True)
                doc.add_paragraph()
                continue
            para = doc.add_heading(block.text, level=min(block.level, 3))
            if block.level == 1 and not first:
                para.paragraph_format.page_break_before = True
            first = False
            set_paragraph_font(para, size=16 if block.level == 1 else 14)
            if block.text == "Mục lục":
                toc_p = doc.add_paragraph()
                add_toc(toc_p)
                doc.add_paragraph()
        elif block.kind == "p":
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            para.paragraph_format.first_line_indent = Cm(0.75)
            para.paragraph_format.line_spacing = 1.15
            run = para.add_run(block.text)
            set_run_font(run, size=13)
        elif block.kind == "bullets":
            for item in block.items:
                para = doc.add_paragraph(style="List Bullet")
                run = para.add_run(item)
                set_run_font(run, size=13)
        elif block.kind == "table" and block.table:
            add_docx_table(doc, block.table)
        elif block.kind == "figure" and block.figure:
            figure = block.figure
            if figure.path.exists():
                p_fig = doc.add_paragraph()
                p_fig.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p_fig.add_run()
                run.add_picture(str(figure.path), width=Inches(5.8))
                cap = doc.add_paragraph()
                cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                r = cap.add_run(figure.caption)
                set_run_font(r, size=11, italic=True)

    return doc


def main() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    blocks = build_blocks()
    OUT_MD.write_text(render_markdown(blocks), encoding="utf-8")
    doc = build_docx(blocks)
    doc.core_properties.title = "Few-Shot Open-Set Keyword Spotting"
    doc.core_properties.subject = "Vietnamese thesis draft"
    doc.core_properties.author = "Codex"
    doc.save(OUT_DOCX)
    print(f"Wrote {OUT_MD}")
    print(f"Wrote {OUT_DOCX}")


if __name__ == "__main__":
    main()
